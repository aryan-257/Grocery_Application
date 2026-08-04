from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '2E7D32')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    return p

def body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_borders(table)
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1B5E20')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = 'F1F8E9' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return table

def fr_table(rows):
    add_table(headers=['ID', 'Requirement'], rows=rows, col_widths=[1.2, 5.3])


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SOFTWARE REQUIREMENTS SPECIFICATION')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('FreshMart — Online Grocery Delivery Web Application')
r.bold = True
r.font.size = Pt(16)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub2.add_run('Microservices Architecture — ASP.NET Core 10 + Angular 21 + Docker')
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
add_table(
    headers=['Field', 'Details'],
    rows=[
        ['Document Version', '2.1'],
        ['Prepared By',      'Software Architecture & Product Team'],
        ['Date',             'April 2026'],
        ['Status',           'Final Draft'],
        ['Architecture',     'Microservices — 6 independent ASP.NET Core 10 services + YARP API Gateway'],
        ['Technology Stack', 'ASP.NET Core 10 | Angular 21 | SQL Server 2022 | SignalR | Razorpay | Docker Compose'],
        ['Classification',   'Confidential — Internal Use Only'],
    ],
    col_widths=[2.0, 4.5]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading('1. Introduction', 1)
heading('1.1 Purpose', 2)
body('This Software Requirements Specification (SRS) document defines all functional and non-functional requirements for FreshMart, a full-featured online grocery delivery web application. FreshMart has been re-architected from a monolithic ASP.NET Core application to a microservices-based platform. The backend is composed of 6 independent ASP.NET Core 10 microservices orchestrated via a YARP API Gateway, with Angular 21 as the Single-Page Application (SPA) frontend. This document serves as the authoritative reference for developers, architects, QA engineers, and all project stakeholders throughout the development lifecycle.')

heading('1.2 Scope', 2)
body('FreshMart enables customers to browse a grocery product catalog organized into 8 departments, build a persistent shopping cart with budget tracking, apply coupon codes, and place orders paid via Razorpay in Indian Rupees (INR). Store Managers manage product inventory, pricing, and discounts. Delivery Drivers view and update assigned orders. Platform Administrators oversee all users, products, orders, coupons, and support tickets. The system delivers real-time order status notifications via ASP.NET Core SignalR, transactional emails via MailKit/SMTP, and a fully integrated customer support ticketing system with live chat. The entire application is containerized using Docker and Docker Compose, with each microservice running in its own container with an isolated SQL Server database.')

heading('1.3 Definitions, Acronyms & Abbreviations', 2)
add_table(
    headers=['Term', 'Definition'],
    rows=[
        ['SRS',          'Software Requirements Specification'],
        ['API',          'Application Programming Interface'],
        ['REST',         'Representational State Transfer — architectural style for HTTP-based web services'],
        ['JWT',          'JSON Web Token — compact, signed token for stateless authentication'],
        ['SPA',          'Single Page Application — the Angular 21 browser-based frontend'],
        ['EF Core',      'Entity Framework Core — .NET ORM used for all database interactions'],
        ['DTO',          'Data Transfer Object — shape of data exchanged between services and Angular client'],
        ['SignalR',      'ASP.NET Core real-time hub library for push notifications and live chat'],
        ['RBAC',         'Role-Based Access Control — permission model based on user roles'],
        ['SKU',          'Stock Keeping Unit — unique identifier for a product'],
        ['OOS',          'Out of Stock — a product with StockQuantity = 0'],
        ['INR',          'Indian Rupee — base currency used throughout FreshMart'],
        ['Razorpay',     'Indian payment gateway supporting cards, UPI, netbanking, and wallets'],
        ['SMTP',         'Simple Mail Transfer Protocol — used via MailKit for transactional emails'],
        ['HMAC',         'Hash-based Message Authentication Code — used for Razorpay signature verification'],
        ['YARP',         'Yet Another Reverse Proxy — Microsoft library used for the API Gateway'],
        ['Microservice', 'An independently deployable service responsible for a single bounded domain'],
        ['API Gateway',  'Single entry point that routes all client requests to appropriate microservices'],
        ['WCAG',         'Web Content Accessibility Guidelines'],
        ['CI/CD',        'Continuous Integration / Continuous Deployment pipeline'],
    ],
    col_widths=[1.3, 5.2]
)

heading('1.4 References', 2)
for ref in [
    'ASP.NET Core 10 Official Documentation — https://learn.microsoft.com/aspnet/core',
    'Angular 21 Official Documentation — https://angular.dev',
    'Entity Framework Core — https://learn.microsoft.com/ef/core',
    'YARP Reverse Proxy Documentation — https://microsoft.github.io/reverse-proxy/',
    'Razorpay API Documentation — https://razorpay.com/docs',
    'MailKit Documentation — https://github.com/jstedfast/MailKit',
    'IEEE Std 830-1998: Recommended Practice for Software Requirements Specifications',
    'OWASP Top 10 Web Application Security Risks',
    'WCAG 2.1 Web Content Accessibility Guidelines',
    'Docker Documentation — https://docs.docker.com',
]:
    bullet(ref)

heading('1.5 Document Overview', 2)
add_table(
    headers=['Section', 'Content'],
    rows=[
        ['Section 2', 'Overall product description, user roles, and operating environment'],
        ['Section 3', 'All functional requirements by module'],
        ['Section 4', 'Ten detailed use cases covering primary workflows'],
        ['Section 5', 'Non-functional requirements across performance, security, reliability'],
        ['Section 6', 'Microservices architecture and technology stack'],
        ['Section 7', 'Per-service data models — entities and relationships'],
        ['Section 8', 'Complete REST API endpoint reference per microservice'],
        ['Section 9', 'Constraints, assumptions, and future enhancement backlog'],
    ],
    col_widths=[1.3, 5.2]
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 2. OVERALL DESCRIPTION
# ══════════════════════════════════════════════════════════════════════════════
heading('2. Overall Description', 1)
heading('2.1 Product Perspective', 2)
body('FreshMart is an independent, containerized web platform providing customers with a seamless online grocery shopping experience. The platform has been migrated from a monolithic architecture to a microservices architecture. Each domain (Auth, Products, Orders, Payments, Notifications, Support) is now an independent ASP.NET Core 10 service with its own dedicated SQL Server 2022 database. All client traffic flows through a YARP-based API Gateway that handles JWT validation and routes requests to the appropriate service. Customers browse 50+ products across 8 categories, add items to a persistent cart, apply coupon codes, and pay via Razorpay in INR. The platform delivers real-time order status notifications through SignalR and HTML transactional emails through MailKit/SMTP.')

heading('2.2 Product Functions — High Level', 2)
add_table(
    headers=['Module', 'Owning Service', 'Description'],
    rows=[
        ['User Management',         'AuthService (port 5001)',         'Registration, login, Google OAuth2, JWT auth, role management (Customer/StoreManager/DeliveryDriver/Admin), profile update, password change'],
        ['Product Catalog',         'ProductService (port 5002)',      '8-category product listings with search, filters, sorting, pagination, discount pricing, stock management, autocomplete, and reviews'],
        ['Shopping Cart',           'OrderService (port 5003)',        'Persistent server-side cart with budget limit tracking, real-time subtotal, discount price calculation, and over-budget alert'],
        ['Order Management',        'OrderService (port 5003)',        'Checkout with coupon validation, delivery fee + tax calculation, Razorpay integration, order history, and status tracking'],
        ['Payment Processing',      'PaymentService (port 5004)',      'Razorpay order creation, HMAC signature verification, webhook handling, payment status tracking — all in INR'],
        ['Coupon System',           'OrderService (port 5003)',        'Percentage and fixed discount coupons with usage limits, minimum order amounts, and expiry dates'],
        ['Real-Time Notifications', 'NotificationService (port 5005)', 'SignalR push notifications for order updates, support replies, new orders (Admin/StoreManager); email via MailKit/SMTP'],
        ['Customer Support',        'SupportService (port 5006)',      'Ticketing system with categories, priorities, status workflow, and real-time live chat via SignalR SupportHub'],
        ['API Gateway',             'ApiGateway (port 8080)',          'YARP reverse proxy routing all client requests to correct microservice; centralized JWT validation and CORS'],
    ],
    col_widths=[1.6, 1.8, 3.1]
)

heading('2.3 User Classes and Characteristics', 2)
add_table(
    headers=['Role', 'Description', 'Key Capabilities'],
    rows=[
        ['Customer',        'Primary end user; browses, shops, and receives deliveries',   'Browse catalog, manage cart, checkout, pay, track orders, write reviews, raise support tickets'],
        ['Store Manager',   'Manages product inventory and order fulfillment',             'Product CRUD, stock management, discount management, order status updates, support ticket management'],
        ['Delivery Driver', 'Picks up and delivers orders to customers',                  'View assigned orders (Shipped/OutForDelivery/Delivered only), update delivery status'],
        ['Platform Admin',  'Oversees all platform operations',                           'Full user/product/order/coupon control, platform statistics, support ticket resolution'],
    ],
    col_widths=[1.3, 2.2, 3.0]
)

heading('2.4 Operating Environment', 2)
add_table(
    headers=['Component', 'Details'],
    rows=[
        ['Frontend',         'Angular 21 SPA with Tailwind CSS; served via Nginx on port 80; Chrome 100+, Firefox 100+, Safari 15+, Edge 100+; mobile-responsive'],
        ['API Gateway',      'ASP.NET Core 10 + YARP on port 8080; routes all /api/v1/* and /hubs/* traffic; validates JWT centrally'],
        ['AuthService',      'ASP.NET Core 10 on port 5001; SQL Server FreshMart_Auth database; BCrypt + JWT + Google OAuth2'],
        ['ProductService',   'ASP.NET Core 10 on port 5002; SQL Server FreshMart_Product database; product catalog, categories, reviews'],
        ['OrderService',     'ASP.NET Core 10 on port 5003; SQL Server FreshMart_Order database; cart, orders, coupons'],
        ['PaymentService',   'ASP.NET Core 10 on port 5004; SQL Server FreshMart_Payment database; Razorpay integration'],
        ['NotificationService','ASP.NET Core 10 on port 5005; SQL Server FreshMart_Notification database; SignalR NotificationHub + MailKit email'],
        ['SupportService',   'ASP.NET Core 10 on port 5006; SQL Server FreshMart_Support database; SignalR SupportHub'],
        ['Database',         'Microsoft SQL Server 2022 — single instance, 6 isolated databases (one per service)'],
        ['Containerization', 'Docker + Docker Compose; all 8 services + SQL Server run via docker compose up --build'],
    ],
    col_widths=[1.6, 4.9]
)

heading('2.5 Design and Implementation Constraints', 2)
for c in [
    'All API endpoints versioned under /api/v1/ with forward-compatible versioning strategy',
    'JWT access token: 1-hour expiry; refresh token: 7-day expiry, stored server-side in AuthService AppUser record',
    'JWT secret key shared across all services via Docker environment variable Jwt__Key for stateless verification',
    'Each microservice has its own isolated SQL Server database — no cross-service direct DB access',
    'Inter-service communication uses typed HttpClient (synchronous HTTP) — no message broker in current version',
    'EF Core Code-First with EnsureCreated() on startup; no manual schema modifications in production',
    'All payment amounts in INR; no card data stored on FreshMart servers (Razorpay handles all card data)',
    'Delivery fee: FREE if SubTotal >= Rs.500, otherwise Rs.49',
    'Tax: 5% of SubTotal applied at order creation',
    'Product discount stored as DiscountPercent (0-100); DiscountedPrice = Price x (1 - DiscountPercent/100)',
    'SignalR hubs require JWT via query string access_token for WebSocket connections',
    'Database auto-created and seeded on startup via service-specific seeders',
    'CORS configured at API Gateway level; allowed origins set via environment variable',
]:
    bullet(c)

heading('2.6 Assumptions and Dependencies', 2)
for a in [
    'Razorpay account configured with Key ID and Key Secret (set via Docker environment variables)',
    'Gmail SMTP account with App Password configured for MailKit email delivery in NotificationService',
    'Docker Desktop installed for local development and container orchestration',
    'Google Cloud project with OAuth2 credentials configured for Google login in AuthService',
    'Product images hosted on external URLs (Unsplash CDN used in seeded data)',
    'All monetary values stored and displayed in Indian Rupees (Rs.)',
    'SQL Server 2022 container starts with health check before any service connects',
    'All services share the same JWT signing key for stateless token verification',
]:
    bullet(a)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 3. FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
heading('3. Functional Requirements', 1)

heading('3.1 User Authentication & Authorization (AuthService)', 2)
fr_table([
    ['FR-AUTH-01', 'The system shall allow users to register with FirstName, LastName, Email, Password, and PhoneNumber; default role is Customer'],
    ['FR-AUTH-02', 'The system shall authenticate via email/password and return a signed JWT (1-hour expiry) plus a refresh token (7-day expiry) stored server-side in AuthService'],
    ['FR-AUTH-03', 'The system shall support Google OAuth2 login via Google userinfo endpoint; auto-create account on first login with GoogleId linked to AppUser record'],
    ['FR-AUTH-04', 'The system shall enforce RBAC for all protected routes: Customer, StoreManager, DeliveryDriver, Admin'],
    ['FR-AUTH-05', 'The system shall support refresh token rotation: POST /api/v1/auth/refresh issues new JWT and new refresh token, invalidating the old one'],
    ['FR-AUTH-06', 'Authenticated users shall update their profile (FirstName, LastName, PhoneNumber) and receive a re-issued JWT with updated claims'],
    ['FR-AUTH-07', 'Authenticated users shall change their password after verifying the current password via POST /api/v1/auth/change-password'],
    ['FR-AUTH-08', 'Admins shall activate or deactivate any user account via PATCH /api/v1/users/{id}/toggle-active'],
    ['FR-AUTH-09', 'Admins shall change any user role (Customer/StoreManager/DeliveryDriver/Admin) via PATCH /api/v1/users/{id}/role'],
    ['FR-AUTH-10', 'The system shall return the current user profile (Id, Email, FirstName, LastName, Role, PhoneNumber) via GET /api/v1/auth/me'],
    ['FR-AUTH-11', 'JWT validation is performed at the API Gateway (YARP) level; individual services also validate JWT independently for defense in depth'],
])

heading('3.2 Product Catalog Management (ProductService)', 2)
fr_table([
    ['FR-CAT-01', 'StoreManagers and Admins shall create products with: Name, Description, Price (INR), SKU, ImageUrl, CategoryId, StockQuantity, Brand, Unit, DiscountPercent'],
    ['FR-CAT-02', 'The system shall support category management via GET /api/v1/categories returning all categories with ParentCategoryId for hierarchical display'],
    ['FR-CAT-03', 'StoreManagers shall update product stock quantities via PATCH /api/v1/products/{id}/stock'],
    ['FR-CAT-04', 'StoreManagers shall set product-level discounts (0-100%) via PATCH /api/v1/products/{id}/discount'],
    ['FR-CAT-05', 'Product listings shall display: Name, Brand, Price, DiscountPercent, DiscountedPrice, Unit, StockQuantity, AverageRating, CategoryName, ImageUrl, IsActive'],
    ['FR-CAT-06', 'Admins shall soft-delete products via DELETE /api/v1/products/{id} (sets IsActive=false); hidden from all public endpoints'],
    ['FR-CAT-07', 'GET /api/v1/products/on-sale returns all products with DiscountPercent > 0, ordered by discount descending'],
    ['FR-CAT-08', 'GET /api/v1/products/low-stock (Admin/StoreManager only) returns products with StockQuantity < 10'],
    ['FR-CAT-09', 'GET /api/v1/products/suggestions?q= returns up to 6 autocomplete matches by name, brand, or category'],
    ['FR-CAT-10', 'The system provides 8 seeded categories: Fruits & Vegetables, Dairy & Eggs, Bakery, Beverages, Snacks, Meat & Seafood, Frozen Foods, Pantry'],
])

heading('3.3 Search & Product Discovery (ProductService)', 2)
fr_table([
    ['FR-SRCH-01', 'Full-text product search across Name, Description, Brand, SKU, and Category Name via GET /api/v1/products?query={q}'],
    ['FR-SRCH-02', 'Search results filterable by: categoryId, minPrice, maxPrice'],
    ['FR-SRCH-03', 'Sorting supported: price_asc, price_desc, rating (default: name ascending)'],
    ['FR-SRCH-04', 'Paginated results with page and pageSize parameters; response includes total count, current page, and page size'],
    ['FR-SRCH-05', 'Only IsActive=true products returned in all public catalog endpoints'],
])

heading('3.4 Shopping Cart (OrderService)', 2)
fr_table([
    ['FR-CART-01', 'The system shall maintain one persistent server-side cart per authenticated customer in OrderService; auto-created on first item add'],
    ['FR-CART-02', 'Cart response includes: Items (ProductId, Name, UnitPrice, ImageUrl, Quantity, TotalPrice, DiscountPercent, OriginalPrice), SubTotal, BudgetLimit, isOverBudget, TotalItems, LastUpdated'],
    ['FR-CART-03', 'Customers set budget limit via PUT /api/v1/cart/budget; isOverBudget=true when SubTotal > BudgetLimit'],
    ['FR-CART-04', 'System validates stock when adding items; returns 400 Insufficient Stock if quantity exceeds StockQuantity'],
    ['FR-CART-05', 'Cart item UnitPrice reflects the discounted price (DiscountPercent applied at time of adding)'],
    ['FR-CART-06', 'Setting quantity=0 in PUT /api/v1/cart/items/{productId} removes the item from cart'],
    ['FR-CART-07', 'Customers clear entire cart via DELETE /api/v1/cart'],
    ['FR-CART-08', 'Customers remove individual items via DELETE /api/v1/cart/items/{productId}'],
])

heading('3.5 Order Management (OrderService)', 2)
fr_table([
    ['FR-ORD-01', 'Customers create orders from cart; system validates cart is not empty before creating order'],
    ['FR-ORD-02', 'System calculates: SubTotal, DeliveryFee (Rs.0 if SubTotal >= Rs.500 else Rs.49), TaxAmount (5% of SubTotal), DiscountAmount (from coupon), TotalAmount'],
    ['FR-ORD-03', 'Customers provide DeliveryAddress and optional Notes at checkout'],
    ['FR-ORD-04', 'System validates coupon at order creation: IsActive, not expired, UsedCount < UsageLimit, SubTotal >= MinOrderAmount'],
    ['FR-ORD-05', 'On order creation, OrderService calls PaymentService via HTTP to create Razorpay order; returns RazorpayOrderId and RazorpayKeyId to frontend'],
    ['FR-ORD-06', 'Order status transitions: Pending -> Processing -> Shipped -> OutForDelivery -> Delivered (or Cancelled at any pre-delivery stage)'],
    ['FR-ORD-07', 'POST /api/v1/orders/{id}/complete-payment clears cart, calls NotificationService to send SignalR notifications and confirmation email'],
    ['FR-ORD-08', 'Each status change via PATCH /api/v1/orders/{id}/status triggers NotificationService call for SignalR notification + transactional email to customer'],
    ['FR-ORD-09', 'DeliveryDriver role sees only orders with status Shipped, OutForDelivery, or Delivered'],
    ['FR-ORD-10', 'EstimatedDelivery set to 2 business days from order creation date'],
    ['FR-ORD-11', 'DeliveredAt timestamp recorded when status set to Delivered'],
])

heading('3.6 Payment Processing (PaymentService)', 2)
fr_table([
    ['FR-PAY-01', 'System integrates Razorpay for all payment processing; no card data stored on FreshMart servers'],
    ['FR-PAY-02', 'POST /api/v1/payment/create-order returns RazorpayOrderId, RazorpayKeyId, Amount, Currency (INR)'],
    ['FR-PAY-03', 'POST /api/v1/payment/verify verifies HMAC SHA256 signature (RazorpayOrderId + "|" + RazorpayPaymentId signed with KeySecret)'],
    ['FR-PAY-04', 'Payment statuses: Pending, Paid, Failed, Refunded, Cancelled'],
    ['FR-PAY-05', 'POST /api/v1/payment/webhook handles Razorpay events authenticated by X-Razorpay-Signature header (no JWT required)'],
    ['FR-PAY-06', 'GET /api/v1/payment/my-payments returns all payments for the current authenticated user'],
    ['FR-PAY-07', 'Payment status lookup by internal PaymentId or by RazorpayOrderId'],
    ['FR-PAY-08', 'All transactions logged: Amount, Currency, Status, RazorpayOrderId, RazorpayPaymentId, PaymentMethod, FailureReason, CreatedAt, CompletedAt'],
    ['FR-PAY-09', 'PaymentService exposes OrderServiceClient to allow OrderService to create Razorpay orders on behalf of customers'],
])

heading('3.7 Coupon System (OrderService)', 2)
fr_table([
    ['FR-CPN-01', 'Admins create coupons: Code, DiscountType (Percentage/Fixed), DiscountValue, MinOrderAmount, ExpiresAt, UsageLimit'],
    ['FR-CPN-02', 'GET /api/v1/coupons returns all active, non-expired coupons (public endpoint)'],
    ['FR-CPN-03', 'POST /api/v1/coupons/validate validates code against order amount; returns validity, type, value, calculated discount amount'],
    ['FR-CPN-04', 'Percentage discount = round(OrderAmount x DiscountValue / 100, 2)'],
    ['FR-CPN-05', 'Fixed discount = min(DiscountValue, OrderAmount) to prevent negative totals'],
    ['FR-CPN-06', 'UsedCount incremented at order creation when valid coupon applied'],
    ['FR-CPN-07', 'Seeded coupons: WELCOME10 (10% off Rs.200+), SAVE50 (Rs.50 off Rs.500+), FRESH20 (20% off Rs.300+), FLAT100 (Rs.100 off Rs.800+), NEWUSER15 (15% off Rs.250+)'],
])

heading('3.8 Reviews & Ratings (ProductService)', 2)
fr_table([
    ['FR-REV-01', 'Customers rate products 1-5 stars with text comment after placing an order containing that product'],
    ['FR-REV-02', 'One review per product per customer; second attempt returns 409 Conflict'],
    ['FR-REV-03', 'Only customers with a non-cancelled order containing the product may review (verified purchase enforcement via OrderProjection in ProductService DB)'],
    ['FR-REV-04', 'Product AverageRating recalculated as arithmetic mean after each new review submission'],
    ['FR-REV-05', 'GET /api/v1/products/{productId}/reviews/can-review returns canReview and alreadyReviewed flags'],
    ['FR-REV-06', 'Reviews ordered by CreatedAt descending; response includes ReviewId, ProductId, CustomerId, CustomerName, Rating, Comment, CreatedAt'],
])

heading('3.9 Real-Time Notifications (NotificationService)', 2)
fr_table([
    ['FR-NOTIF-01', 'NotificationHub delivers real-time push for: order status changes, payment confirmation, support replies, new orders (Admin/StoreManager)'],
    ['FR-NOTIF-02', 'On SignalR connection, user automatically joins group user:{userId} and role:{role} for targeted and broadcast notifications'],
    ['FR-NOTIF-03', 'Notifications persisted to FreshMart_Notification DB: Title, Message, Type (info/success/warning/error/order), Link, IsRead, CreatedAt'],
    ['FR-NOTIF-04', 'GET /api/v1/notifications returns last 50 notifications ordered by CreatedAt descending'],
    ['FR-NOTIF-05', 'PATCH /api/v1/notifications/{id}/read marks single notification read; PATCH /api/v1/notifications/read-all marks all read'],
    ['FR-NOTIF-06', 'GET /api/v1/notifications/unread-count returns count of unread notifications'],
    ['FR-NOTIF-07', 'DELETE /api/v1/notifications/{id} deletes single; DELETE /api/v1/notifications deletes all for user'],
    ['FR-NOTIF-08', 'Admin and StoreManager receive real-time SignalR notifications for new orders and new support tickets'],
    ['FR-NOTIF-09', 'NotificationService exposes internal HTTP endpoint for other services (OrderService, SupportService) to trigger notifications'],
    ['FR-NOTIF-10', 'HTML transactional emails sent via MailKit/SMTP for all 6 order lifecycle events; failures logged but do not block API response'],
])

heading('3.10 Customer Support (SupportService)', 2)
fr_table([
    ['FR-SUP-01', 'Customers create support tickets with: Subject, Category (Order/Payment/Delivery/Product/Other), Description (initial message), Priority (Low/Medium/High)'],
    ['FR-SUP-02', 'Ticket status workflow: Open -> InProgress -> Resolved -> Closed'],
    ['FR-SUP-03', 'When staff (Admin/StoreManager) first replies to an Open ticket, system auto-moves status to InProgress'],
    ['FR-SUP-04', 'Customers see only their own tickets; Admin and StoreManager see all tickets with filters for status, priority, and category'],
    ['FR-SUP-05', 'Both customers and staff add messages to ticket thread via POST /api/v1/support/tickets/{id}/messages'],
    ['FR-SUP-06', 'New messages broadcast in real time via SignalR SupportHub to all participants in ticket:{ticketId} group'],
    ['FR-SUP-07', 'Admin/StoreManager update ticket status and priority via PATCH /api/v1/support/tickets/{id}/status'],
    ['FR-SUP-08', 'SupportService maintains its own AppUser projection (synced from AuthService JWT claims) for sender name resolution'],
])

heading('3.11 Admin Dashboard', 2)
fr_table([
    ['FR-ADM-01', 'Admins view platform statistics: total users, active/inactive counts, users by role breakdown (via AuthService /api/v1/users/stats)'],
    ['FR-ADM-02', 'Admins manage all users: list with filters (role, search, isActive), view detail, update profile, change role, toggle active, delete'],
    ['FR-ADM-03', 'Admins manage all products: create, update, soft-delete, update stock, update discount (via ProductService)'],
    ['FR-ADM-04', 'Admins manage all orders: view all orders across all customers, update order status (via OrderService)'],
    ['FR-ADM-05', 'Admins manage coupons: create new coupons, view all active coupons (via OrderService)'],
    ['FR-ADM-06', 'Admins manage all support tickets: view, reply, update status and priority (via SupportService)'],
    ['FR-ADM-07', 'Admins receive real-time SignalR notifications for new orders and new support tickets (via NotificationService)'],
])
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 4. DETAILED USE CASES
# ══════════════════════════════════════════════════════════════════════════════
heading('4. Detailed Use Cases', 1)
body('This section presents ten comprehensive use cases covering the primary workflows of the FreshMart platform. Each use case specifies actors, preconditions, complete step-by-step normal and alternative flows, exception handling, business rules, and the microservice(s) involved.')

# UC-01
heading('UC-01: Customer Registration & Login', 2)
add_table(headers=['Field', 'Details'], rows=[
    ['Use Case ID',    'UC-01'],
    ['Use Case Name',  'Customer Registration and Login'],
    ['Actor(s)',       'New Customer, Returning Customer'],
    ['Service',        'AuthService (port 5001) via API Gateway (port 8080)'],
    ['Description',    'A new visitor registers a FreshMart account using email/password or Google OAuth2, then logs in to access the platform.'],
    ['Preconditions',  'User is not logged in. User has a valid email address.'],
    ['Postconditions', 'Account created. JWT + refresh token issued. User lands on home page.'],
    ['Trigger',        'User clicks Sign Up or Login from the home page or when accessing a protected feature.'],
], col_widths=[1.5, 5.0])
body('Normal Flow:')
for s in [
    '1. User navigates to /register',
    '2. User enters: FirstName, LastName, Email, Password, PhoneNumber',
    '3. Angular validates fields in real time: required fields, email format, password strength',
    '4. User clicks Create Account',
    '5. Angular POSTs to POST /api/v1/auth/register (routed by Gateway to AuthService)',
    '6. AuthService checks email uniqueness — returns 409 if duplicate',
    '7. AuthService BCrypt-hashes password and creates AppUser (Role=Customer, IsActive=true)',
    '8. AuthService returns { userId, email, role }',
    '9. Angular redirects to /login with success banner',
    '10. User enters credentials; Angular POSTs to POST /api/v1/auth/login',
    '11. AuthService returns { accessToken, refreshToken, expiresAt, role, userId }',
    '12. Angular stores token; redirects to /home',
]:
    bullet(s)
body('Alternative Flow — Google OAuth2:')
for s in [
    '1. User clicks Continue with Google',
    '2. Angular initiates Google OAuth2 flow; gets ID token',
    '3. Angular POSTs to POST /api/v1/auth/google with { idToken }',
    '4. AuthService validates token via Google userinfo endpoint (https://www.googleapis.com/oauth2/v3/userinfo)',
    '5. If new user: auto-create account with Role=Customer; if existing: link GoogleId',
    '6. JWT issued; user lands on /home',
]:
    bullet(s)
body('Exception Flow:')
for s in [
    'Duplicate email -> 409: Email already registered',
    'Invalid Google token -> 401: Invalid Google token',
    'Wrong password -> 401: Invalid credentials',
]:
    bullet(s)
body('API Endpoints: POST /api/v1/auth/register | POST /api/v1/auth/login | POST /api/v1/auth/google | POST /api/v1/auth/refresh | POST /api/v1/auth/logout')

# UC-02
heading('UC-02: Browse & Search Products', 2)
add_table(headers=['Field', 'Details'], rows=[
    ['Use Case ID',    'UC-02'],
    ['Use Case Name',  'Browse and Search Products with Filters'],
    ['Actor(s)',       'Customer, Guest User'],
    ['Service',        'ProductService (port 5002) via API Gateway'],
    ['Description',    'A user searches for grocery products, applies filters and sorting, and adds items to the cart.'],
    ['Preconditions',  'At least one active product exists in the catalog.'],
    ['Postconditions', 'User found desired products; authenticated users can add to cart.'],
    ['Trigger',        'User types in search bar, clicks a category, or navigates to /products.'],
], col_widths=[1.5, 5.0])
body('Normal Flow:')
for s in [
    '1. Customer types "Amul Milk" in search bar',
    '2. Angular debounces 300ms; GET /api/v1/products?query=Amul+Milk&page=1&pageSize=20',
    '3. API Gateway routes to ProductService; searches Name, Description, Brand, SKU, Category.Name (IsActive=true only)',
    '4. Returns paginated result: { items[], total, page, pageSize }',
    '5. Angular renders product grid: Name, Brand, Price, DiscountedPrice, DiscountPercent badge, StockQuantity, AverageRating, Unit, ImageUrl',
    '6. Customer filters: categoryId=Dairy & Eggs, minPrice=50, maxPrice=300',
    '7. Customer sorts by price_asc',
    '8. Customer clicks Add to Cart on Amul Milk 1L (Rs.65)',
    '9. POST /api/v1/cart/items { productId, quantity: 1 } — routed to OrderService',
    '10. OrderService validates stock; adds to cart; returns updated cart',
    '11. Angular updates cart icon count and shows toast notification',
]:
    bullet(s)
body('API Endpoints: GET /api/v1/products | GET /api/v1/products/on-sale | GET /api/v1/products/suggestions | GET /api/v1/categories | POST /api/v1/cart/items')

# UC-03
heading('UC-03: Shopping Cart & Budget Tracking', 2)
add_table(headers=['Field', 'Details'], rows=[
    ['Use Case ID',    'UC-03'],
    ['Use Case Name',  'Shopping Cart Management and Budget Tracking'],
    ['Actor(s)',       'Customer'],
    ['Service',        'OrderService (port 5003) via API Gateway'],
    ['Description',    'Customer manages their cart, sets a budget limit, and monitors spending in real time.'],
    ['Preconditions',  'Customer is authenticated.'],
    ['Postconditions', 'Cart updated with correct prices, quantities, and budget status.'],
    ['Trigger',        'Customer navigates to /cart or adds an item from the product catalog.'],
], col_widths=[1.5, 5.0])
body('Normal Flow:')
for s in [
    '1. Customer navigates to /cart',
    '2. GET /api/v1/cart (OrderService) returns items with discounted prices, SubTotal, BudgetLimit, isOverBudget',
    '3. Cart shows: Amul Milk Rs.65 x 2 = Rs.130, Britannia Bread Rs.35 x 1 = Rs.35, SubTotal = Rs.165',
    '4. Customer sets budget: PUT /api/v1/cart/budget { budgetLimit: 500 }',
    '5. Budget bar shows: Rs.165 of Rs.500 (green)',
    '6. Customer adds Mutton 500g (Rs.450); SubTotal = Rs.615',
    '7. isOverBudget = true; budget bar turns red: Rs.615 — Rs.115 over budget',
    '8. Customer removes Mutton: DELETE /api/v1/cart/items/{productId}',
    '9. SubTotal returns to Rs.165; isOverBudget = false',
]:
    bullet(s)
body('API Endpoints: GET /api/v1/cart | POST /api/v1/cart/items | PUT /api/v1/cart/items/{productId} | DELETE /api/v1/cart/items/{productId} | DELETE /api/v1/cart | PUT /api/v1/cart/budget')

# UC-04
heading('UC-04: Checkout, Coupon & Order Placement', 2)
add_table(headers=['Field', 'Details'], rows=[
    ['Use Case ID',    'UC-04'],
    ['Use Case Name',  'Checkout with Coupon and Razorpay Payment'],
    ['Actor(s)',       'Customer'],
    ['Services',       'OrderService -> PaymentService -> NotificationService (inter-service HTTP calls)'],
    ['Description',    'Customer proceeds through checkout: reviews cart, enters delivery address, applies coupon, pays via Razorpay, and places order.'],
    ['Preconditions',  'Customer authenticated. Cart has at least one item.'],
    ['Postconditions', 'Order placed, Razorpay order created, cart cleared after payment, confirmation email sent.'],
    ['Trigger',        'Customer clicks Proceed to Checkout.'],
], col_widths=[1.5, 5.0])
body('Normal Flow:')
for s in [
    '1. Customer reviews cart: SubTotal Rs.850',
    '2. Customer enters DeliveryAddress: 14 MG Road, Pune 411001',
    '3. Customer enters coupon code FRESH20',
    '4. POST /api/v1/coupons/validate { code: FRESH20, orderAmount: 850 } -> OrderService',
    '5. OrderService returns: { valid: true, discountType: Percentage, discountValue: 20, discountAmount: Rs.170 }',
    '6. Order summary: SubTotal Rs.850, Delivery Rs.0 (>=Rs.500), Tax Rs.42.50 (5%), Discount -Rs.170, Total Rs.722.50',
    '7. Customer clicks Place Order',
    '8. POST /api/v1/orders { deliveryAddress, notes, couponCode: FRESH20 } -> OrderService',
    '9. OrderService creates Order record; calls PaymentService HTTP to create Razorpay order',
    '10. Returns { order, razorpayOrderId, razorpayKey }',
    '11. Angular initializes Razorpay checkout with razorpayOrderId',
    '12. Customer completes payment (UPI/card/netbanking)',
    '13. Razorpay returns { razorpay_order_id, razorpay_payment_id, razorpay_signature }',
    '14. POST /api/v1/payment/verify -> PaymentService verifies HMAC signature',
    '15. POST /api/v1/orders/{id}/complete-payment -> OrderService clears cart, calls NotificationService for notifications + email',
    '16. Angular redirects to /orders/{id} showing order confirmation',
]:
    bullet(s)
body('API Endpoints: POST /api/v1/orders | POST /api/v1/coupons/validate | POST /api/v1/payment/create-order | POST /api/v1/payment/verify | POST /api/v1/orders/{id}/complete-payment')

# UC-05
heading('UC-05: Razorpay Payment & Webhook Handling', 2)
add_table(headers=['Field', 'Details'], rows=[
    ['Use Case ID',    'UC-05'],
    ['Use Case Name',  'Razorpay Payment Processing and Webhook Handling'],
    ['Actor(s)',       'Customer, Razorpay System'],
    ['Service',        'PaymentService (port 5004) via API Gateway'],
    ['Description',    'Payment is processed via Razorpay; webhook events update payment status automatically.'],
    ['Preconditions',  'Order created with RazorpayOrderId. Customer on payment page.'],
    ['Postconditions', 'Payment recorded as Paid; order status updated; customer notified.'],
    ['Trigger',        'Customer initiates payment on Razorpay checkout.'],
], col_widths=[1.5, 5.0])
body('Normal Flow:')
for s in [
    '1. Angular loads Razorpay checkout with { key, amount, orderId, name: FreshMart }',
    '2. Customer selects UPI; enters UPI ID',
    '3. Razorpay processes payment; returns { razorpay_order_id, razorpay_payment_id, razorpay_signature }',
    '4. POST /api/v1/payment/verify { razorpayOrderId, razorpayPaymentId, razorpaySignature } -> PaymentService',
    '5. PaymentService computes HMAC SHA256: hmac(razorpayOrderId + "|" + razorpayPaymentId, KeySecret)',
    '6. Signature matches -> Payment.Status = Paid; Payment.CompletedAt = now',
    '7. POST /api/v1/orders/{id}/complete-payment -> OrderService',
    '8. Cart cleared; SignalR notification sent via NotificationService; confirmation email sent',
]:
    bullet(s)
body('Webhook Flow:')
for s in [
    '1. Razorpay sends POST to /api/v1/payment/webhook with X-Razorpay-Signature header',
    '2. PaymentService verifies webhook signature',
    '3. payment.captured event -> Payment.Status = Paid',
    '4. payment.failed event -> Payment.Status = Failed; FailureReason recorded',
    '5. refund.created event -> Payment.Status = Refunded',
]:
    bullet(s)
body('API Endpoints: POST /api/v1/payment/create-order | POST /api/v1/payment/verify | POST /api/v1/payment/webhook | GET /api/v1/payment/{id}/status | GET /api/v1/payment/my-payments')

# UC-06
heading('UC-06: Order Status Tracking & Notifications', 2)
add_table(headers=['Field', 'Details'], rows=[
    ['Use Case ID',    'UC-06'],
    ['Use Case Name',  'Order Status Tracking and Notifications'],
    ['Actor(s)',       'Customer, Store Manager, Delivery Driver, Admin'],
    ['Services',       'OrderService -> NotificationService (HTTP call for each status update)'],
    ['Description',    'After order placement, staff update order status. Each update triggers a real-time SignalR notification and transactional email to the customer via NotificationService.'],
    ['Preconditions',  'Order exists with status Pending or later. Staff member is authenticated.'],
    ['Postconditions', 'Order status updated. Customer notified via SignalR and email.'],
    ['Trigger',        'Store Manager or Admin clicks Update Status on an order.'],
], col_widths=[1.5, 5.0])
body('Normal Flow:')
for s in [
    '1. Store Manager navigates to /admin/orders',
    '2. GET /api/v1/orders (OrderService) returns all orders for Admin/StoreManager',
    '3. Manager clicks Update Status -> Processing',
    '4. PATCH /api/v1/orders/{id}/status { status: Processing } -> OrderService',
    '5. OrderService updates Order.Status = Processing',
    '6. OrderService calls NotificationService HTTP: sends SignalR notification to customer',
    '7. OrderService calls AuthService HTTP to get customer email; calls NotificationService to send email',
    '8. Customer receives real-time SignalR push and email for each status change',
    '9. Delivery Driver updates to OutForDelivery -> Delivered; DeliveredAt timestamp recorded',
]:
    bullet(s)
body('API Endpoints: GET /api/v1/orders | GET /api/v1/orders/{id} | PATCH /api/v1/orders/{id}/status')

# UC-07
heading('UC-07: Product Reviews & Ratings', 2)
add_table(headers=['Field', 'Details'], rows=[
    ['Use Case ID',    'UC-07'],
    ['Use Case Name',  'Product Reviews and Ratings'],
    ['Actor(s)',       'Customer'],
    ['Service',        'ProductService (port 5002) — uses OrderProjection table for purchase verification'],
    ['Description',    'After receiving a delivered order, the customer rates and reviews individual products. Only verified purchasers can review.'],
    ['Preconditions',  'Customer has a non-cancelled order containing the product. Customer has not yet reviewed this product.'],
    ['Postconditions', 'Review published. Product AverageRating updated.'],
    ['Trigger',        'Customer navigates to product detail page and clicks Write a Review.'],
], col_widths=[1.5, 5.0])
body('Normal Flow:')
for s in [
    '1. Customer navigates to /products/{id}',
    '2. GET /api/v1/products/{productId}/reviews/can-review returns { canReview: true, alreadyReviewed: false }',
    '3. ProductService checks OrderProjection table (synced from OrderService) for verified purchase',
    '4. Customer selects 4 stars and types review comment',
    '5. POST /api/v1/products/{productId}/reviews { rating: 4, comment: "..." }',
    '6. ProductService validates purchase; creates Review record',
    '7. ProductService recalculates AverageRating = average of all ratings for this product',
    '8. Review appears on product detail page ordered by newest first',
]:
    bullet(s)
body('API Endpoints: GET /api/v1/products/{productId}/reviews | GET /api/v1/products/{productId}/reviews/can-review | POST /api/v1/products/{productId}/reviews')

# UC-08
heading('UC-08: Customer Support Ticketing & Live Chat', 2)
add_table(headers=['Field', 'Details'], rows=[
    ['Use Case ID',    'UC-08'],
    ['Use Case Name',  'Customer Support Ticketing and Live Chat'],
    ['Actor(s)',       'Customer, Store Manager, Admin'],
    ['Service',        'SupportService (port 5006) — SignalR SupportHub for real-time chat'],
    ['Description',    'Customer raises a support ticket. Staff respond via real-time live chat via SignalR SupportHub.'],
    ['Preconditions',  'Customer is authenticated.'],
    ['Postconditions', 'Ticket created. Messages exchanged in real time. Ticket resolved and closed.'],
    ['Trigger',        'Customer clicks Contact Support or Help.'],
], col_widths=[1.5, 5.0])
body('Normal Flow:')
for s in [
    '1. Customer creates ticket: POST /api/v1/support/tickets { subject, category, description, priority }',
    '2. SupportService creates SupportTicket (Status=Open) and initial SupportMessage',
    '3. Admin opens ticket; Angular connects to SupportHub: JoinTicket(ticketId)',
    '4. Admin types reply; POST /api/v1/support/tickets/{id}/messages',
    '5. SupportService creates SupportMessage (IsStaff=true); auto-moves ticket to InProgress',
    '6. SupportService broadcasts via SignalR SupportHub to ticket:{ticketId} group: newMessage event',
    '7. Customer sees reply in real time without page refresh',
    '8. Admin resolves: PATCH /api/v1/support/tickets/{id}/status { status: Resolved }',
    '9. SupportHub broadcasts ticketUpdated event to all participants',
]:
    bullet(s)
body('API Endpoints: POST /api/v1/support/tickets | GET /api/v1/support/tickets | GET /api/v1/support/tickets/{id} | POST /api/v1/support/tickets/{id}/messages | PATCH /api/v1/support/tickets/{id}/status')

# UC-09
heading('UC-09: Admin User & Platform Management', 2)
add_table(headers=['Field', 'Details'], rows=[
    ['Use Case ID',    'UC-09'],
    ['Use Case Name',  'Admin User and Platform Management'],
    ['Actor(s)',       'Platform Admin'],
    ['Service',        'AuthService (port 5001) for user management'],
    ['Description',    'Admin manages all platform users, views statistics, changes roles, and activates/deactivates accounts.'],
    ['Preconditions',  'Admin is authenticated with Admin role.'],
    ['Postconditions', 'User accounts managed. Platform statistics reviewed.'],
    ['Trigger',        'Admin navigates to /admin/users or /admin/dashboard.'],
], col_widths=[1.5, 5.0])
body('Normal Flow:')
for s in [
    '1. Admin navigates to /admin/dashboard',
    '2. GET /api/v1/users/stats (AuthService) returns { total, active, inactive, byRole }',
    '3. Admin navigates to /admin/users; GET /api/v1/users?role=Customer&isActive=true',
    '4. Admin changes role: PATCH /api/v1/users/{id}/role { role: StoreManager }',
    '5. Admin deactivates user: PATCH /api/v1/users/{id}/toggle-active',
    '6. User can no longer log in (IsActive=false check in AuthService login)',
]:
    bullet(s)
body('API Endpoints: GET /api/v1/users | GET /api/v1/users/stats | GET /api/v1/users/{id} | PUT /api/v1/users/{id} | PATCH /api/v1/users/{id}/role | PATCH /api/v1/users/{id}/toggle-active | DELETE /api/v1/users/{id}')

# UC-10
heading('UC-10: Store Manager Inventory & Discount Management', 2)
add_table(headers=['Field', 'Details'], rows=[
    ['Use Case ID',    'UC-10'],
    ['Use Case Name',  'Store Manager Inventory and Discount Management'],
    ['Actor(s)',       'Store Manager'],
    ['Service',        'ProductService (port 5002)'],
    ['Description',    'Store Manager manages product inventory, updates stock levels, sets discounts, and monitors low-stock products.'],
    ['Preconditions',  'Store Manager is authenticated with StoreManager role.'],
    ['Postconditions', 'Product catalog updated. Discounts applied. Low-stock products identified.'],
    ['Trigger',        'Store Manager navigates to /admin/products or receives low-stock alert.'],
], col_widths=[1.5, 5.0])
body('Normal Flow:')
for s in [
    '1. Store Manager navigates to /admin/products',
    '2. GET /api/v1/products (ProductService) returns all products with stock levels',
    '3. GET /api/v1/products/low-stock returns all products with StockQuantity < 10',
    '4. Manager updates stock: PATCH /api/v1/products/{id}/stock { quantity: 100 }',
    '5. Manager sets discount: PATCH /api/v1/products/{id}/discount { discountPercent: 15 }',
    '6. Product now shows 15% OFF badge; DiscountedPrice = Price x 0.85',
    '7. Manager creates new product: POST /api/v1/products { name, description, price, sku, ... }',
]:
    bullet(s)
body('API Endpoints: GET /api/v1/products | GET /api/v1/products/low-stock | POST /api/v1/products | PUT /api/v1/products/{id} | PATCH /api/v1/products/{id}/stock | PATCH /api/v1/products/{id}/discount | DELETE /api/v1/products/{id}')
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. NON-FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
heading('5. Non-Functional Requirements', 1)

heading('5.1 Performance', 2)
add_table(
    headers=['ID', 'Requirement', 'Target'],
    rows=[
        ['NFR-PERF-01', 'Product catalog and search API response time (ProductService)',    '< 200ms (p95)'],
        ['NFR-PERF-02', 'Cart operations (add, remove, update) via OrderService',           '< 300ms (p95)'],
        ['NFR-PERF-03', 'Order placement end-to-end latency (OrderService + PaymentService HTTP call)', '< 2s (p95)'],
        ['NFR-PERF-04', 'Angular SPA Largest Contentful Paint (LCP)',                       '< 2.5s on 4G mobile connection'],
        ['NFR-PERF-05', 'SignalR notification delivery latency (NotificationService)',       '< 100ms per notification to connected clients'],
        ['NFR-PERF-06', 'API Gateway routing overhead',                                     '< 5ms additional latency per request'],
        ['NFR-PERF-07', 'Database query time for order history (paginated)',                 '< 100ms with proper composite indexing'],
        ['NFR-PERF-08', 'Autocomplete suggestions response time (ProductService)',           '< 150ms for up to 6 results'],
    ],
    col_widths=[1.2, 3.0, 2.3]
)

heading('5.2 Security', 2)
for s in [
    'JWT validation performed at API Gateway (YARP) level for all incoming requests',
    'Each microservice independently validates JWT as a second layer of defense',
    'All services share the same Jwt__Key, Jwt__Issuer, Jwt__Audience via Docker environment variables',
    'HTTPS enforced at all entry points; HTTP requests permanently redirected to HTTPS',
    'Passwords hashed with BCrypt (cost factor 12) in AuthService; never stored in plain text',
    'SQL injection prevention via EF Core parameterized queries and LINQ in all services',
    'XSS prevention: Angular built-in DOM sanitizer on all user-generated content',
    'CORS restricted to registered Angular origin at API Gateway level',
    'Razorpay HMAC SHA256 signature verification on all payment confirmations in PaymentService',
    'Razorpay webhook signature verified via X-Razorpay-Signature header',
    'No card data stored on FreshMart servers; all card handling via Razorpay',
    'JWT access token: 1-hour expiry; refresh token: 7-day expiry with server-side rotation in AuthService',
    'Inter-service HTTP calls use JWT forwarding (Bearer token passed from gateway to services)',
    'RabbitMQ used for async inter-service events via MassTransit (OrderPlacedEvent, PaymentCompletedEvent, OrderStatusChangedEvent)',
    'MassTransit retry policy: 1s, 5s, 15s intervals before dead-lettering failed messages',
]:
    bullet(s)

heading('5.3 Reliability & Availability', 2)
add_table(
    headers=['ID', 'Requirement', 'Target'],
    rows=[
        ['NFR-REL-01', 'Platform uptime SLA',                    '99.9% (< 8.7 hours downtime per year)'],
        ['NFR-REL-02', 'Order data durability',                  'Zero data loss; all order writes within EF Core transactions in OrderService'],
        ['NFR-REL-03', 'Database backup',                        'Daily full backup + hourly incremental; 30-day retention per service database'],
        ['NFR-REL-04', 'Recovery Time Objective (RTO)',          '< 1 hour'],
        ['NFR-REL-05', 'Recovery Point Objective (RPO)',         '< 15 minutes'],
        ['NFR-REL-06', 'SignalR reconnection',                   'Automatic reconnection with exponential backoff on client disconnect'],
        ['NFR-REL-07', 'Email delivery failure handling',        'Failures logged in NotificationService; do not block API response'],
        ['NFR-REL-08', 'Razorpay webhook retry handling',        'Idempotent webhook processing in PaymentService; duplicate events safely ignored'],
        ['NFR-REL-09', 'Service startup order',                  'Docker Compose healthcheck ensures SQL Server is ready before any service starts'],
        ['NFR-REL-10', 'Inter-service HTTP failure handling',    'Non-critical service calls (email, notifications) wrapped in try/catch; do not fail primary operation'],
    ],
    col_widths=[1.2, 2.8, 2.5]
)

heading('5.4 Scalability', 2)
for s in [
    'Stateless API design (JWT) enables horizontal scaling of any individual microservice independently',
    'Each microservice has its own database — no shared state between services',
    'Angular SPA deployed to CDN for global edge distribution',
    'Docker Compose for local development; Kubernetes-ready for production scaling',
    'Database connection pooling via EF Core for efficient resource utilization per service',
    'SignalR supports scale-out via Azure SignalR Service or Redis backplane for multi-server deployment',
    'API Gateway (YARP) can be scaled independently from backend services',
    'Individual services can be scaled based on load (e.g., ProductService scaled more than SupportService)',
]:
    bullet(s)

heading('5.5 Usability & Accessibility', 2)
for s in [
    'Angular 21 with Tailwind CSS for consistent, responsive UI across all screen sizes',
    'Mobile-first responsive layout tested across 320px (small mobile) to 2560px (wide desktop)',
    'Real-time form validation with clear error messages on all input fields',
    'Toast notifications for all user actions (add to cart, order placed, payment confirmed)',
    'Loading states and skeleton screens for all async data fetching operations',
]:
    bullet(s)

heading('5.6 Maintainability & DevOps', 2)
for s in [
    'Each microservice is independently deployable — changes to one service do not require redeploying others',
    'Each service has its own Dockerfile using multi-stage builds (SDK build stage + ASP.NET runtime stage)',
    'Docker Compose orchestrates all 8 containers (6 services + API Gateway + SQL Server + Frontend) with one command',
    'Environment-specific configuration via appsettings.json and Docker environment variables per service',
    'Structured logging via ILogger throughout all controllers and services',
    'OpenAPI/Swagger documentation auto-generated per service; available at :{port}/swagger',
    'SharedModels project for shared DTOs and event contracts between services',
    'Angular lazy-loaded feature modules for optimal bundle size and load performance',
]:
    bullet(s)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 6. SYSTEM ARCHITECTURE & TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
heading('6. System Architecture & Technology Stack', 1)

heading('6.1 Architecture Overview', 2)
body('FreshMart uses a microservices architecture. The Angular 21 SPA communicates exclusively with the YARP API Gateway over HTTP/HTTPS. The Gateway validates JWT tokens and routes requests to the appropriate microservice based on URL path. Each microservice owns its domain logic and its own SQL Server database. Services communicate with each other via synchronous HTTP calls using typed HttpClient. SignalR hubs in NotificationService and SupportService handle all bidirectional real-time events. The entire application is containerized using Docker and Docker Compose.')

heading('6.2 Microservices Architecture Diagram', 2)
body('Request Flow: Browser -> Nginx (port 80) -> Angular SPA -> API Gateway (port 8080) -> Microservices (ports 5001-5006) -> SQL Server (port 1433)')
body('Synchronous HTTP Calls: OrderService -> PaymentService (create Razorpay order) | OrderService -> NotificationService (send notifications/emails) | OrderService -> AuthService (get customer email for status emails)')
body('Async RabbitMQ Events (MassTransit): OrderService publishes OrderPlacedEvent -> consumed by ProductService (stock decrement) and NotificationService (email + push). PaymentService publishes PaymentCompletedEvent -> consumed by OrderService (transition to Processing). OrderService publishes OrderStatusChangedEvent -> consumed by NotificationService.')

heading('6.3 Full Technology Stack', 2)
add_table(
    headers=['Layer / Concern', 'Technology', 'Purpose'],
    rows=[
        ['Frontend SPA',          'Angular 21 + Tailwind CSS',                    'UI, routing, state management, HTTP communication'],
        ['Web Server',            'Nginx (Alpine)',                                'Serves Angular static files; proxies API calls'],
        ['API Gateway',           'ASP.NET Core 10 + YARP',                       'Reverse proxy; JWT validation; CORS; routes to microservices'],
        ['Auth Microservice',     'ASP.NET Core 10 Web API',                      'User registration, login, JWT issuance, Google OAuth2, user management'],
        ['Product Microservice',  'ASP.NET Core 10 Web API',                      'Product catalog, categories, search, reviews, stock management'],
        ['Order Microservice',    'ASP.NET Core 10 Web API',                      'Cart, orders, coupons, order status management'],
        ['Payment Microservice',  'ASP.NET Core 10 Web API',                      'Razorpay integration, payment verification, webhook handling'],
        ['Notification Microservice','ASP.NET Core 10 + SignalR + MailKit',       'Real-time push notifications, transactional emails'],
        ['Support Microservice',  'ASP.NET Core 10 + SignalR',                    'Support ticketing, live chat via SupportHub'],
        ['ORM & Migrations',      'Entity Framework Core (Code-First)',            'All database read/write per service; EnsureCreated on startup'],
        ['Database',              'Microsoft SQL Server 2022',                     '6 isolated databases: FreshMart_Auth, _Product, _Order, _Payment, _Notification, _Support'],
        ['Authentication',        'JWT Bearer + Refresh Tokens + Google OAuth2',  'Stateless auth; social login; shared secret across all services'],
        ['Payment',               'Razorpay API (HTTP) + Webhooks',               'Card/UPI/netbanking payments in INR'],
        ['Email',                 'MailKit + Gmail SMTP',                          'All transactional emails (order lifecycle events)'],
        ['Message Broker',        'RabbitMQ + MassTransit',                       'Async inter-service events (OrderPlaced, PaymentCompleted, OrderStatusChanged)'],
        ['Containerization',      'Docker + Docker Compose',                       'One-command deployment; all services containerized'],
        ['Shared Models',         'Microservices/SharedModels project',            'Shared DTOs and event contracts between services'],
    ],
    col_widths=[1.7, 2.1, 2.7]
)

heading('6.4 API Gateway Routing Configuration', 2)
body('The YARP API Gateway (appsettings.json) routes all traffic based on URL path prefix:')
add_table(
    headers=['Route', 'Path Pattern', 'Target Service', 'Port'],
    rows=[
        ['auth-route',          '/api/v1/auth/**',          'AuthService',         '5001'],
        ['users-route',         '/api/v1/users/**',         'AuthService',         '5001'],
        ['products-route',      '/api/v1/products/**',      'ProductService',      '5002'],
        ['categories-route',    '/api/v1/categories/**',    'ProductService',      '5002'],
        ['orders-route',        '/api/v1/orders/**',        'OrderService',        '5003'],
        ['cart-route',          '/api/v1/cart/**',          'OrderService',        '5003'],
        ['coupons-route',       '/api/v1/coupons/**',       'OrderService',        '5003'],
        ['payment-route',       '/api/v1/payment/**',       'PaymentService',      '5004'],
        ['notifications-route', '/api/v1/notifications/**', 'NotificationService', '5005'],
        ['notif-hub-route',     '/hubs/notifications/**',   'NotificationService', '5005'],
        ['support-route',       '/api/v1/support/**',       'SupportService',      '5006'],
        ['support-hub-route',   '/hubs/support/**',         'SupportService',      '5006'],
    ],
    col_widths=[1.5, 1.8, 1.7, 0.7]
)

heading('6.5 Docker Deployment Architecture', 2)
add_table(
    headers=['Container', 'Port', 'Database', 'Purpose'],
    rows=[
        ['freshmart-sqlserver',      '1433',      'N/A',                      'SQL Server 2022 — shared instance, 6 isolated databases'],
        ['freshmart-gateway',        '8080',      'None',                     'YARP API Gateway — routes all client traffic'],
        ['freshmart-auth',           '5001',      'FreshMart_Auth',           'AuthService — users, JWT, Google OAuth2'],
        ['freshmart-products',       '5002',      'FreshMart_Product',        'ProductService — catalog, categories, reviews'],
        ['freshmart-orders',         '5003',      'FreshMart_Order',          'OrderService — cart, orders, coupons'],
        ['freshmart-payment',        '5004',      'FreshMart_Payment',        'PaymentService — Razorpay integration'],
        ['freshmart-notifications',  '5005',      'FreshMart_Notification',   'NotificationService — SignalR + email'],
        ['freshmart-support',        '5006',      'FreshMart_Support',        'SupportService — tickets + live chat'],
        ['freshmart-frontend',       '80',        'None',                     'Angular 21 SPA served via Nginx'],
    ],
    col_widths=[1.7, 0.6, 1.6, 2.6]
)

heading('6.6 Frontend Angular Structure', 2)
for s in [
    'src/app/core/guards/ — AuthGuard, RoleGuard for route protection',
    'src/app/core/interceptors/ — AuthInterceptor (attaches JWT to all HTTP requests)',
    'src/app/core/models/ — TypeScript interfaces for all DTOs',
    'src/app/core/services/ — AuthService, CartService, ProductService, OrderService, PaymentService, NotificationService, WishlistService, ComparisonService, InvoiceService, RecentlyViewedService, ThemeService, LocationService, CouponService',
    'src/app/shared/components/ — Navbar, ProductCard, SearchBar (reusable components)',
    'src/app/pages/home/ — Home page with featured products and categories',
    'src/app/pages/products/ — Product listing, search, filters, product detail',
    'src/app/pages/cart/ — Cart page with budget tracker',
    'src/app/pages/checkout/ — Checkout flow with Razorpay integration',
    'src/app/pages/orders/ — Order history and order detail',
    'src/app/pages/order-tracking/ — Real-time order tracking',
    'src/app/pages/profile/ — User profile, password change',
    'src/app/pages/admin/ — Admin dashboard: users, products, orders, support',
    'src/app/pages/store-manager/ — Store manager dashboard',
    'src/app/pages/support/ — Customer support ticket creation and chat',
    'src/app/pages/offers/ — On-sale products page',
    'src/app/pages/delivery/ — Delivery driver view',
    'src/app/pages/compare/ — Product comparison page',
    'src/app/pages/rate-order/ — Order rating page',
]:
    bullet(s)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 7. DATA MODEL — PER SERVICE
# ══════════════════════════════════════════════════════════════════════════════
heading('7. Data Model — Per Microservice', 1)
body('Each microservice owns its own database and entity models. There are no cross-service foreign keys. Services that need data from another service either call that service via HTTP or maintain a local projection/cache of the required data.')

heading('7.1 AuthService — FreshMart_Auth Database', 2)
add_table(
    headers=['Entity', 'Key Attributes', 'Notes'],
    rows=[
        ['AppUser', 'Id (Guid PK), Email, PasswordHash (BCrypt), FirstName, LastName, Role (Customer/StoreManager/DeliveryDriver/Admin), PhoneNumber, IsActive, RefreshToken, RefreshTokenExpiry, GoogleId, CreatedAt',
         'Central user store; JWT claims sourced from this entity; RefreshToken rotated on every refresh call'],
    ],
    col_widths=[1.3, 3.5, 1.7]
)

heading('7.2 ProductService — FreshMart_Product Database', 2)
add_table(
    headers=['Entity', 'Key Attributes', 'Notes'],
    rows=[
        ['Category',         'Id (Guid PK), Name, Description, ImageUrl, ParentCategoryId (FK, nullable, self-referencing)',
                             'Self-referencing hierarchy; 8 seeded categories'],
        ['Product',          'Id (Guid PK), Name, Description, Price (decimal INR), Sku, ImageUrl, CategoryId (FK), StockQuantity, IsActive, AverageRating, Brand, Unit, DiscountPercent (0-100), CreatedAt',
                             'DiscountedPrice computed: Price x (1 - DiscountPercent/100)'],
        ['Review',           'Id (Guid PK), ProductId (FK), CustomerId (Guid), CustomerName (snapshot), Rating (1-5), Comment, CreatedAt',
                             'One per customer per product; AverageRating recalculated on each new review'],
        ['AppUser',          'Id (Guid PK), Email, FirstName, LastName (projection)',
                             'Local projection of AuthService user for review author name resolution'],
        ['OrderProjection',  'Id (Guid PK), CustomerId (Guid), Status',
                             'Local projection of orders for verified-purchase review enforcement'],
        ['OrderItemProjection','Id (Guid PK), OrderId (FK), ProductId (FK)',
                             'Local projection of order items for verified-purchase check'],
    ],
    col_widths=[1.5, 3.2, 1.8]
)

heading('7.3 OrderService — FreshMart_Order Database', 2)
add_table(
    headers=['Entity', 'Key Attributes', 'Notes'],
    rows=[
        ['Cart',     'Id (Guid PK), CustomerId (Guid), BudgetLimit (decimal, nullable), LastUpdated',
                     'One cart per customer; auto-created on first item add'],
        ['CartItem', 'Id (Guid PK), CartId (FK), ProductId (Guid), Quantity',
                     'Belongs to Cart; ProductId references ProductService (no FK constraint)'],
        ['Order',    'Id (Guid PK), CustomerId (Guid), Status (Pending/Processing/Shipped/OutForDelivery/Delivered/Cancelled), SubTotal, DeliveryFee, TaxAmount, DiscountAmount, TotalAmount, DeliveryAddress, Notes, CreatedAt, EstimatedDelivery, DeliveredAt',
                     'DeliveryFee = Rs.0 if SubTotal >= Rs.500 else Rs.49; Tax = 5% of SubTotal'],
        ['OrderItem','Id (Guid PK), OrderId (FK), ProductId (Guid), ProductName (snapshot), Quantity, UnitPrice',
                     'ProductName snapshotted at order time to preserve history'],
        ['Coupon',   'Id (Guid PK), Code, DiscountType (Percentage/Fixed), DiscountValue, MinOrderAmount, ExpiresAt (nullable), IsActive, UsageLimit, UsedCount',
                     '5 coupons seeded on startup'],
        ['Product',  'Id (Guid PK), Name, Price, DiscountPercent (local cache)',
                     'Local product cache in OrderService for price calculation without cross-service calls'],
    ],
    col_widths=[1.3, 3.4, 1.8]
)

heading('7.4 PaymentService — FreshMart_Payment Database', 2)
add_table(
    headers=['Entity', 'Key Attributes', 'Notes'],
    rows=[
        ['Payment', 'Id (Guid PK), UserId (Guid), OrderId (Guid), Amount, Currency (INR), RazorpayOrderId, RazorpayPaymentId, RazorpaySignature, Status (Pending/Paid/Failed/Refunded/Cancelled), PaymentMethod, FailureReason, CreatedAt, CompletedAt, Metadata',
                    'All Razorpay transaction data stored here; no card data stored'],
    ],
    col_widths=[1.3, 3.5, 1.7]
)

heading('7.5 NotificationService — FreshMart_Notification Database', 2)
add_table(
    headers=['Entity', 'Key Attributes', 'Notes'],
    rows=[
        ['Notification', 'Id (Guid PK), UserId (Guid), Title, Message, Type (info/success/warning/error/order), Link, IsRead, CreatedAt',
                         'Last 50 returned per user; real-time delivery via SignalR NotificationHub'],
        ['AppUser',      'Id (Guid PK), Email, FirstName, LastName (projection)',
                         'Local projection for notification targeting'],
    ],
    col_widths=[1.3, 3.5, 1.7]
)

heading('7.6 SupportService — FreshMart_Support Database', 2)
add_table(
    headers=['Entity', 'Key Attributes', 'Notes'],
    rows=[
        ['SupportTicket',  'Id (Guid PK), CustomerId (Guid), CustomerName, CustomerEmail, Subject, Category (Order/Payment/Delivery/Product/Other), Status (Open/InProgress/Resolved/Closed), Priority (Low/Medium/High), CreatedAt, UpdatedAt, ResolvedAt',
                           'Status auto-moves to InProgress on first staff reply'],
        ['SupportMessage', 'Id (Guid PK), TicketId (FK), SenderId (Guid), SenderName, SenderRole, Message, IsStaff, CreatedAt',
                           'Broadcast via SignalR SupportHub to ticket:{ticketId} group on creation'],
        ['AppUser',        'Id (Guid PK), Email, FirstName, LastName (projection)',
                           'Local projection for sender name resolution'],
    ],
    col_widths=[1.3, 3.5, 1.7]
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 8. API ENDPOINTS REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
heading('8. API Endpoints Reference', 1)
body('All endpoints are accessed via the API Gateway at port 8080. The Gateway routes to the appropriate microservice based on URL path. Protected endpoints require a valid JWT Bearer token in the Authorization header. Role annotations indicate the minimum role required.')

heading('8.1 AuthService Endpoints (/api/v1/auth, /api/v1/users)', 2)
add_table(
    headers=['Method', 'Endpoint', 'Description', 'Auth'],
    rows=[
        ['POST', '/api/v1/auth/register',         'Register new user (FirstName, LastName, Email, Password, PhoneNumber)',  'None'],
        ['POST', '/api/v1/auth/login',             'Login — returns JWT access token + refresh token',                      'None'],
        ['POST', '/api/v1/auth/refresh',           'Rotate refresh token and issue new JWT',                               'Refresh Token'],
        ['POST', '/api/v1/auth/logout',            'Revoke refresh token',                                                 'JWT'],
        ['GET',  '/api/v1/auth/me',                'Get current user profile',                                             'JWT'],
        ['PUT',  '/api/v1/auth/me',                'Update profile (FirstName, LastName, PhoneNumber)',                    'JWT'],
        ['POST', '/api/v1/auth/change-password',   'Change password (currentPassword, newPassword)',                       'JWT'],
        ['POST', '/api/v1/auth/google',            'Google OAuth2 login/register via ID token',                            'None'],
        ['GET',  '/api/v1/users',                  'List all users with filters (role, search, isActive)',                 'Admin'],
        ['GET',  '/api/v1/users/stats',            'Platform user statistics (total, active, by role)',                    'Admin'],
        ['GET',  '/api/v1/users/{id}',             'Get user by ID',                                                       'Admin'],
        ['PUT',  '/api/v1/users/{id}',             'Update user profile',                                                  'Admin'],
        ['PATCH','/api/v1/users/{id}/role',        'Change user role',                                                     'Admin'],
        ['PATCH','/api/v1/users/{id}/toggle-active','Activate or deactivate user account',                                 'Admin'],
        ['DELETE','/api/v1/users/{id}',            'Permanently delete user',                                              'Admin'],
    ],
    col_widths=[0.7, 2.3, 2.8, 0.7]
)

heading('8.2 ProductService Endpoints (/api/v1/products, /api/v1/categories)', 2)
add_table(
    headers=['Method', 'Endpoint', 'Description', 'Auth'],
    rows=[
        ['GET',   '/api/v1/products',                    'List products with filters (query, categoryId, minPrice, maxPrice, sortBy, page, pageSize)', 'None'],
        ['GET',   '/api/v1/products/suggestions',        'Autocomplete suggestions (q param, max 6 results)',                                          'None'],
        ['GET',   '/api/v1/products/on-sale',            'Products with DiscountPercent > 0, ordered by discount desc',                               'None'],
        ['GET',   '/api/v1/products/low-stock',          'Products with StockQuantity < 10',                                                          'Admin/Manager'],
        ['GET',   '/api/v1/products/{id}',               'Get full product detail',                                                                   'None'],
        ['POST',  '/api/v1/products',                    'Create new product listing',                                                                'Admin/Manager'],
        ['PUT',   '/api/v1/products/{id}',               'Update product details, price, stock',                                                      'Admin/Manager'],
        ['DELETE','/api/v1/products/{id}',               'Soft-delete product (IsActive=false)',                                                      'Admin'],
        ['PATCH', '/api/v1/products/{id}/stock',         'Update stock quantity only',                                                                'Admin/Manager'],
        ['PATCH', '/api/v1/products/{id}/discount',      'Set discount percent (0-100)',                                                              'Admin/Manager'],
        ['GET',   '/api/v1/categories',                  'Get all categories with ParentCategoryId',                                                  'None'],
        ['GET',   '/api/v1/products/{id}/reviews',       'Get paginated product reviews',                                                             'None'],
        ['GET',   '/api/v1/products/{id}/reviews/can-review', 'Check if current user can review this product',                                       'JWT'],
        ['POST',  '/api/v1/products/{id}/reviews',       'Submit product review (rating, comment)',                                                   'JWT (Customer)'],
    ],
    col_widths=[0.7, 2.5, 2.5, 0.8]
)

heading('8.3 OrderService Endpoints (/api/v1/cart, /api/v1/orders, /api/v1/coupons)', 2)
add_table(
    headers=['Method', 'Endpoint', 'Description', 'Auth'],
    rows=[
        ['GET',    '/api/v1/cart',                    'Get current cart with items, subtotal, budget status',  'JWT'],
        ['POST',   '/api/v1/cart/items',              'Add product to cart (productId, quantity)',             'JWT'],
        ['PUT',    '/api/v1/cart/items/{productId}',  'Update cart item quantity (0 = remove)',                'JWT'],
        ['DELETE', '/api/v1/cart/items/{productId}',  'Remove specific item from cart',                       'JWT'],
        ['DELETE', '/api/v1/cart',                    'Clear entire cart',                                     'JWT'],
        ['PUT',    '/api/v1/cart/budget',             'Set or update cart budget limit',                       'JWT'],
        ['GET',    '/api/v1/orders',                  'Get orders (customer: own; admin/manager: all)',        'JWT'],
        ['GET',    '/api/v1/orders/{id}',             'Get full order detail',                                 'JWT'],
        ['POST',   '/api/v1/orders',                  'Create order from cart (deliveryAddress, notes, couponCode)', 'JWT (Customer)'],
        ['POST',   '/api/v1/orders/{id}/complete-payment', 'Confirm payment, clear cart, send notifications', 'JWT (Customer)'],
        ['PATCH',  '/api/v1/orders/{id}/status',      'Update order status',                                  'Admin/Manager/Driver'],
        ['GET',    '/api/v1/coupons',                 'List all active, non-expired coupons',                 'None'],
        ['POST',   '/api/v1/coupons/validate',        'Validate coupon code against order amount',            'JWT'],
    ],
    col_widths=[0.7, 2.3, 2.8, 0.7]
)

heading('8.4 PaymentService Endpoints (/api/v1/payment)', 2)
add_table(
    headers=['Method', 'Endpoint', 'Description', 'Auth'],
    rows=[
        ['POST',  '/api/v1/payment/create-order',          'Create Razorpay order; return RazorpayOrderId and key',        'JWT'],
        ['POST',  '/api/v1/payment/verify',                'Verify Razorpay HMAC signature after payment',                 'JWT'],
        ['POST',  '/api/v1/payment/webhook',               'Razorpay webhook event receiver',                              'Razorpay-Sig'],
        ['GET',   '/api/v1/payment/{id}/status',           'Get payment status by internal PaymentId',                     'JWT'],
        ['GET',   '/api/v1/payment/order/{razorpayId}/status', 'Get payment status by RazorpayOrderId',                   'JWT'],
        ['GET',   '/api/v1/payment/my-payments',           'Get all payments for current user',                            'JWT'],
    ],
    col_widths=[0.7, 2.5, 2.5, 0.8]
)

heading('8.5 NotificationService Endpoints (/api/v1/notifications)', 2)
add_table(
    headers=['Method', 'Endpoint', 'Description', 'Auth'],
    rows=[
        ['GET',   '/api/v1/notifications',              'Get last 50 notifications for current user',              'JWT'],
        ['GET',   '/api/v1/notifications/unread-count', 'Get unread notification count',                           'JWT'],
        ['PATCH', '/api/v1/notifications/{id}/read',    'Mark single notification as read',                        'JWT'],
        ['PATCH', '/api/v1/notifications/read-all',     'Mark all notifications as read',                          'JWT'],
        ['DELETE','/api/v1/notifications/{id}',         'Delete single notification',                              'JWT'],
        ['DELETE','/api/v1/notifications',              'Delete all notifications for current user',               'JWT'],
    ],
    col_widths=[0.7, 2.3, 2.8, 0.7]
)

heading('8.6 SupportService Endpoints (/api/v1/support)', 2)
add_table(
    headers=['Method', 'Endpoint', 'Description', 'Auth'],
    rows=[
        ['POST',  '/api/v1/support/tickets',               'Create support ticket (subject, category, description, priority)', 'JWT'],
        ['GET',   '/api/v1/support/tickets',               'List tickets (customer: own; admin/manager: all with filters)',    'JWT'],
        ['GET',   '/api/v1/support/tickets/{id}',          'Get ticket detail with all messages',                             'JWT'],
        ['POST',  '/api/v1/support/tickets/{id}/messages', 'Add message to ticket thread',                                    'JWT'],
        ['PATCH', '/api/v1/support/tickets/{id}/status',   'Update ticket status and priority',                               'Admin/Manager'],
    ],
    col_widths=[0.7, 2.5, 2.5, 0.8]
)

heading('8.7 SignalR Hubs', 2)
add_table(
    headers=['Hub', 'Service', 'Endpoint', 'Events / Methods', 'Auth'],
    rows=[
        ['NotificationHub', 'NotificationService', '/hubs/notifications',
         'Server->Client: receiveNotification\nOnConnected: joins user:{userId} and role:{role} groups',
         'JWT (query string access_token)'],
        ['SupportHub', 'SupportService', '/hubs/support',
         'Client->Server: JoinTicket(ticketId), LeaveTicket(ticketId)\nServer->Client: newMessage, ticketUpdated',
         'JWT (query string access_token)'],
    ],
    col_widths=[1.2, 1.3, 1.3, 2.4, 0.8]
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 9. CONSTRAINTS & ASSUMPTIONS
# ══════════════════════════════════════════════════════════════════════════════
heading('9. Constraints & Assumptions', 1)

heading('9.1 Technical Constraints', 2)
add_table(
    headers=['Type', 'Constraint'],
    rows=[
        ['Framework',        'All microservices target .NET 10 (ASP.NET Core 10); Angular version 21 or above required'],
        ['Database',         'SQL Server 2022 for all services; one instance, 6 isolated databases (one per service)'],
        ['Security',         'Passwords must be BCrypt-hashed in AuthService; JWT access token must not be stored in localStorage'],
        ['Payment',          'All payments in INR via Razorpay; no card data stored on FreshMart servers'],
        ['Images',           'Product images hosted on external CDN URLs; no local file upload in v1'],
        ['Cart State',       'Cart must be persisted server-side in OrderService database; one cart per authenticated user'],
        ['Real-Time',        'SignalR requires JWT via query string access_token for WebSocket upgrade'],
        ['Inter-Service',    'Services communicate via synchronous HTTP (typed HttpClient) for critical calls; RabbitMQ + MassTransit for async event-driven communication'],
        ['Browser',          'Angular SPA must support last 2 major versions of Chrome, Firefox, Edge, and Safari'],
        ['Delivery Fee',     'Free delivery for orders >= Rs.500; Rs.49 delivery fee for orders below Rs.500'],
        ['Tax',              '5% tax applied on SubTotal at order creation; not configurable in v1'],
        ['Service Isolation','No direct database access between services; data shared only via HTTP API calls or local projections'],
    ],
    col_widths=[1.3, 5.2]
)

heading('9.2 Business Assumptions', 2)
for a in [
    'FreshMart operates as a single-store platform in v1; multi-store support is a v2 feature',
    'All products are managed by Store Managers and Admins; customer self-listing is not supported',
    'Delivery drivers are pre-vetted and onboarded by the platform operations team',
    'All stores operate within a single country (India) and currency (INR) at v1 launch',
    'Razorpay handles all currency conversion; FreshMart stores amounts in INR',
    'Loyalty points system is a v2 feature; not implemented in v1',
    'Product images are hosted on external CDN; image upload functionality is a v2 feature',
    'Estimated delivery is set to 2 business days; real-time delivery tracking is a v2 feature',
    'Support tickets are handled manually by Admin/StoreManager; AI chatbot is a v2 feature',
    'Coupon codes are created by Admins only; store-specific promotions are a v2 feature',
]:
    bullet(a)

heading('9.3 Seeded Test Data', 2)
body('The following test accounts are seeded on first startup for development and testing:')
add_table(
    headers=['Role', 'Email', 'Password'],
    rows=[
        ['Admin',          'admin@grocery.com',    'Admin@123'],
        ['Store Manager',  'manager@grocery.com',  'Manager@123'],
        ['Delivery Driver','driver@grocery.com',   'Driver@123'],
        ['Customer',       'customer@grocery.com', 'Customer@123'],
    ],
    col_widths=[1.3, 2.5, 1.7]
)
body('Seeded product catalog: 50+ products across 8 categories with Indian brands (Amul, Tata, Britannia, Haldiram, Lays, Maggi, MDH, Fortune, India Gate, Bisleri, Nescafe, Cadbury, McCain, Safal, Suguna, Keggfarms).')
body('Seeded coupons:')
add_table(
    headers=['Code', 'Type', 'Value', 'Min Order', 'Expiry'],
    rows=[
        ['WELCOME10', 'Percentage', '10%',   'Rs.200',  '6 months'],
        ['SAVE50',    'Fixed',      'Rs.50', 'Rs.500',  '3 months'],
        ['FRESH20',   'Percentage', '20%',   'Rs.300',  '2 months'],
        ['FLAT100',   'Fixed',      'Rs.100','Rs.800',  '1 month'],
        ['NEWUSER15', 'Percentage', '15%',   'Rs.250',  '12 months'],
    ],
    col_widths=[1.2, 1.0, 0.8, 1.0, 1.0]
)

heading('9.4 Future Enhancements — Out of Scope for v1.0', 2)
for f in [
    'Asynchronous inter-service communication via RabbitMQ or Azure Service Bus (event-driven architecture)',
    'Service discovery and health monitoring (Consul, Kubernetes liveness/readiness probes)',
    'Distributed tracing across microservices (OpenTelemetry, Jaeger)',
    'Centralized logging aggregation (ELK Stack or Azure Monitor)',
    'Native iOS and Android mobile applications (React Native or Flutter)',
    'Multi-store support: multiple grocery stores under one platform',
    'Real-time delivery tracking with live driver GPS map',
    'Loyalty points and rewards program',
    'Product image upload to cloud storage (Azure Blob / AWS S3)',
    'AI-powered product recommendations based on purchase history',
    'AI chatbot for customer support automation',
    'Subscription box service: curated weekly grocery boxes',
    'Multi-currency and multi-country support',
    'Push notifications via Web Push API / Firebase Cloud Messaging',
    'SMS notifications via Twilio for order dispatch and delivery alerts',
    'Advanced analytics dashboard with charts and KPI metrics',
    'Redis caching layer for product catalog and session data',
    'API rate limiting per service via middleware or API Gateway policies',
]:
    bullet(f)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# REVISION HISTORY
# ══════════════════════════════════════════════════════════════════════════════
heading('Document Revision History', 1)
add_table(
    headers=['Version', 'Date', 'Author', 'Description'],
    rows=[
        ['0.1', 'January 2026',  'Architecture Team', 'Initial draft — monolithic architecture, scope definition, user roles, module overview'],
        ['0.5', 'February 2026', 'Architecture Team', 'Functional requirements, data model, and tech stack added (monolithic)'],
        ['0.9', 'March 2026',    'Architecture Team', '10 detailed use cases, NFRs, full API endpoint reference added (monolithic)'],
        ['1.0', 'March 2026',    'Architecture Team', 'Final monolithic version — ASP.NET Core 10 single API + Angular 19'],
        ['2.1', 'April 2026',    'Architecture Team', 'Added RabbitMQ + MassTransit async messaging (OrderPlacedEvent, PaymentCompletedEvent, OrderStatusChangedEvent); updated inter-service communication section; removed message broker from future enhancements'],
    ],
    col_widths=[0.7, 1.3, 1.8, 2.7]
)

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run('End of Document — FreshMart Grocery Delivery Web Application SRS v2.0 (Microservices Edition)')
r.bold = True
r.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save('FreshMart_SRS.docx')
print('SUCCESS: FreshMart_SRS.docx has been generated!')
