"""
FreshMart — Project Structure & Flow Document Generator (with embedded diagrams)
Run generate_diagrams.py first to create the diagrams/ folder, then run this.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

DIAG = "diagrams"

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '2E7D32')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    return p

def body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_borders(table)
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1B5E20')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = 'F1F8E9' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return table

def insert_image(filename, width_inches=6.5, caption=None):
    path = os.path.join(DIAG, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs:
                r.font.size = Pt(9)
                r.font.italic = True
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()
    else:
        body(f"[Diagram not found: {path} — run generate_diagrams.py first]")


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FRESHMART — PROJECT STRUCTURE & FLOW DOCUMENT')
run.bold = True; run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('File Structure · Architecture · Flowcharts · ER Diagrams · 0-Level DFD')
r.bold = True; r.font.size = Pt(13)

doc.add_paragraph()
add_table(
    headers=['Field', 'Details'],
    rows=[
        ['Project',       'FreshMart — Online Grocery Delivery Platform'],
        ['Architecture',  'Microservices — 6 ASP.NET Core 10 services + YARP API Gateway'],
        ['Frontend',      'Angular 21 + Tailwind CSS + Nginx'],
        ['Database',      'SQL Server 2022 — 6 isolated databases (one per service)'],
        ['Real-Time',     'ASP.NET Core SignalR (NotificationHub + SupportHub)'],
        ['Payments',      'Razorpay (INR) — cards, UPI, netbanking'],
        ['Container',     'Docker + Docker Compose — 9 containers total'],
        ['Document Date', 'April 2026'],
    ],
    col_widths=[1.8, 4.7]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading('1. Project Overview', 1)
body('FreshMart is a full-stack online grocery delivery platform built with a microservices architecture. The project was originally a monolithic ASP.NET Core application and has been re-architected into 6 independent microservices, each owning its own domain and database. All client traffic flows through a YARP-based API Gateway. The Angular 21 frontend is served via Nginx. Everything runs in Docker containers orchestrated by Docker Compose.')
body('The platform supports 4 user roles: Customer (browse, shop, pay, track), Store Manager (manage products and orders), Delivery Driver (update delivery status), and Admin (full platform control). Key features include real-time notifications via SignalR, transactional emails via MailKit, Razorpay payment integration, product reviews, coupon system, and a live support chat.')

heading('1.1 Technology Summary', 2)
add_table(
    headers=['Layer', 'Technology', 'Version', 'Role'],
    rows=[
        ['Frontend',          'Angular',          '21',        'SPA — UI, routing, HTTP calls to API Gateway'],
        ['Frontend Server',   'Nginx',             'Alpine',    'Serves Angular static files on port 80'],
        ['API Gateway',       'YARP + ASP.NET',   '10.0',      'Routes all /api/v1/* traffic; validates JWT centrally'],
        ['Auth Service',      'ASP.NET Core',     '10.0',      'User management, JWT issuance, Google OAuth2'],
        ['Product Service',   'ASP.NET Core',     '10.0',      'Product catalog, categories, reviews, stock'],
        ['Order Service',     'ASP.NET Core',     '10.0',      'Cart, orders, coupons, checkout'],
        ['Payment Service',   'ASP.NET Core',     '10.0',      'Razorpay integration, payment verification'],
        ['Notification Svc',  'ASP.NET Core',     '10.0',      'SignalR push + MailKit email notifications'],
        ['Support Service',   'ASP.NET Core',     '10.0',      'Support tickets + SignalR live chat'],
        ['Database',          'SQL Server',        '2022',      '6 isolated databases, one per service'],
        ['ORM',               'EF Core',           '10.0',      'Code-First, EnsureCreated on startup'],
        ['Auth',              'JWT Bearer',        'HS256',     'Stateless auth; shared secret across all services'],
        ['Payments',          'Razorpay',          'v1 API',    'INR payments — cards, UPI, netbanking, wallets'],
        ['Email',             'MailKit',           '4.8',       'SMTP via Gmail — transactional emails'],
        ['Real-Time',         'SignalR',           '1.2',       'WebSocket push notifications and live chat'],
        ['Message Broker',    'RabbitMQ + MassTransit', '8.3.6', 'Async events: OrderPlaced, PaymentCompleted, OrderStatusChanged'],
        ['Containers',        'Docker Compose',    'v3',        '9 containers: 6 services + gateway + SQL + frontend'],
    ],
    col_widths=[1.4, 1.3, 0.8, 3.0]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE DIAGRAM
# ══════════════════════════════════════════════════════════════════════════════
heading('2. System Architecture', 1)
body('The diagram below shows the complete system architecture of FreshMart. The Angular 21 SPA runs in the browser and communicates only with the API Gateway. The Gateway validates JWT tokens and routes each request to the correct microservice based on URL path. Each microservice has its own isolated SQL Server database. Services communicate with each other via direct HTTP calls. SignalR hubs in NotificationService and SupportService handle real-time WebSocket connections. External integrations include Razorpay for payments, Gmail SMTP for emails, and Google OAuth2 for social login.')
insert_image('01_architecture.png', width_inches=6.8,
             caption='Figure 1: FreshMart System Architecture — Microservices with API Gateway')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. 0-LEVEL DFD (CONTEXT DIAGRAM)
# ══════════════════════════════════════════════════════════════════════════════
heading('3. Level 0 DFD — Context Diagram', 1)
body('The Level 0 DFD (Context Diagram) shows FreshMart as a single system and all external entities that interact with it. It defines the system boundary and shows what data flows in and out. External entities are: Customer, Store Manager, Admin, Delivery Driver, Razorpay (payment gateway), Gmail SMTP (email), and Google OAuth2 (social login).')

heading('3.1 External Entities & Data Flows', 2)
add_table(
    headers=['External Entity', 'Data INTO System', 'Data FROM System'],
    rows=[
        ['Customer',        'Register/Login credentials, Product search queries, Cart items, Order details, Coupon codes, Payment info, Support tickets, Review ratings',
                            'JWT token, Product catalog, Cart state, Order confirmation, Order status updates, Real-time notifications, Invoices, Support replies'],
        ['Store Manager',   'Product create/update data, Stock updates, Discount settings, Order status updates',
                            'Product list, Order list, Low-stock alerts, Real-time notifications for new orders'],
        ['Admin',           'User management actions, Coupon creation, Role changes, Platform queries',
                            'User list, Platform statistics, All orders, All support tickets, All products'],
        ['Delivery Driver', 'Delivery status updates (Shipped/OutForDelivery/Delivered)',
                            'Assigned orders list, Delivery details, Real-time notifications'],
        ['Razorpay',        'Payment status events (webhook), Payment ID + signature after payment',
                            'Create payment order request, Signature verification request'],
        ['Gmail SMTP',      'Email delivery status',
                            'HTML transactional emails for 6 order lifecycle events'],
        ['Google OAuth2',   'User profile (name, email, Google sub ID)',
                            'ID token verification request'],
    ],
    col_widths=[1.4, 2.8, 2.3]
)
insert_image('02_dfd_level0.png', width_inches=6.8,
             caption='Figure 2: Level 0 DFD — FreshMart Context Diagram')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. PROJECT STRUCTURE FLOWCHART
# ══════════════════════════════════════════════════════════════════════════════
heading('4. Project Folder Structure', 1)
body('The workspace root contains two main folders: Frontend (Angular 21 SPA) and Microservices (all backend services). Docker Compose at the root orchestrates everything. The diagram below shows the complete folder hierarchy with all key files.')
insert_image('03_project_structure.png', width_inches=6.8,
             caption='Figure 3: FreshMart Project Folder Structure')

heading('4.1 Root Level Files', 2)
add_table(
    headers=['Path', 'Description'],
    rows=[
        ['docker-compose.yml',      'Orchestrates all 9 containers — SQL Server, API Gateway, 6 microservices, Frontend'],
        ['FreshMart_SRS.docx',      'Software Requirements Specification document (v2.0 — microservices edition)'],
        ['generate_srs.py',         'Python script that generates the SRS DOCX using python-docx'],
        ['generate_project_doc.py', 'Python script that generates this project structure document'],
        ['generate_diagrams.py',    'Python script that generates all architecture/ER/flow diagrams using matplotlib'],
        ['Frontend/',               'Angular 21 SPA — complete frontend application'],
        ['Microservices/',          'All backend microservices — 6 services + API Gateway + SharedModels'],
        ['diagrams/',               'Auto-generated PNG diagram images (created by generate_diagrams.py)'],
    ],
    col_widths=[2.2, 4.3]
)
doc.add_page_break()

heading('4.2 Frontend Structure', 2)
body('The Angular 21 frontend uses standalone components (no NgModules), lazy loading for all routes, and Tailwind CSS for styling. It is built and served via Nginx in Docker.')
add_table(
    headers=['Path', 'Description'],
    rows=[
        ['Frontend/src/main.ts',                          'Angular bootstrap entry point'],
        ['Frontend/src/index.html',                       'Root HTML shell — Angular mounts here'],
        ['Frontend/src/styles.css',                       'Global Tailwind CSS imports and custom styles'],
        ['Frontend/src/app/app.ts',                       'Root AppComponent — contains router-outlet'],
        ['Frontend/src/app/app.routes.ts',                'All 25 application routes with lazy loading and guards'],
        ['Frontend/src/app/app.config.ts',                'Angular app config — registers HttpClient, Router, interceptors'],
        ['Frontend/src/app/core/guards/auth.guard.ts',    'AuthGuard — redirects unauthenticated users to /auth/login'],
        ['Frontend/src/app/core/guards/role.guard.ts',    'RoleGuard — checks JWT role claim; redirects to /unauthorized if mismatch'],
        ['Frontend/src/app/core/interceptors/auth.interceptor.ts', 'Attaches JWT Bearer token to every outgoing API request; handles 401 refresh'],
        ['Frontend/src/app/core/models/index.ts',         'All TypeScript interfaces: User, Product, Cart, Order, Payment, Notification, etc.'],
        ['Frontend/src/app/core/services/auth.service.ts',         'Login, register, Google OAuth2, token refresh, logout, profile'],
        ['Frontend/src/app/core/services/cart.service.ts',         'Cart CRUD, budget management'],
        ['Frontend/src/app/core/services/product.service.ts',      'Product listing, search, filters, suggestions, categories'],
        ['Frontend/src/app/core/services/order.service.ts',        'Order creation, history, status tracking'],
        ['Frontend/src/app/core/services/notification.service.ts', 'Notification CRUD + SignalR connection management'],
        ['Frontend/src/app/shared/components/navbar/',             'Navbar — search, cart icon, notifications bell, user menu'],
        ['Frontend/src/app/shared/components/product-card/',       'Reusable product card component'],
        ['Frontend/src/app/shared/components/search-bar/',         'Search bar with autocomplete suggestions'],
        ['Frontend/src/app/pages/home/home.ts',                    'Home page — featured products, categories, banners'],
        ['Frontend/src/app/pages/products/products.ts',            'Product listing — search, filters, sorting, pagination'],
        ['Frontend/src/app/pages/product-detail/product-detail.ts','Product detail — full info, reviews, add to cart'],
        ['Frontend/src/app/pages/cart/cart.ts',                    'Cart page — items, quantities, budget tracker'],
        ['Frontend/src/app/pages/checkout/checkout.ts',            'Checkout — address, coupon, Razorpay payment'],
        ['Frontend/src/app/pages/orders/orders.ts',                'Order history page'],
        ['Frontend/src/app/pages/order-tracking/order-tracking.ts','Real-time order status tracking'],
        ['Frontend/src/app/pages/rate-order/rate-order.ts',        'Submit star rating and review for ordered products'],
        ['Frontend/src/app/pages/offers/offers.ts',                'All products with active discounts'],
        ['Frontend/src/app/pages/compare/compare.ts',              'Side-by-side product comparison'],
        ['Frontend/src/app/pages/profile/profile.ts',              'Edit name, phone, change password'],
        ['Frontend/src/app/pages/support/support.ts',              'Create ticket, view tickets, live chat via SignalR'],
        ['Frontend/src/app/pages/delivery/delivery.ts',            'Delivery driver — view assigned orders, update status'],
        ['Frontend/src/app/pages/store-manager/manager-dashboard.ts','Store manager dashboard'],
        ['Frontend/src/app/pages/admin/dashboard/',                'Admin dashboard — platform stats'],
        ['Frontend/src/app/pages/admin/products/',                 'Admin product CRUD, stock, discounts'],
        ['Frontend/src/app/pages/admin/orders/',                   'Admin orders — all orders, status updates'],
        ['Frontend/src/app/pages/admin/users/',                    'Admin users — role change, activate/deactivate'],
        ['Frontend/src/app/pages/admin/support/',                  'Admin support — all tickets, reply, status'],
        ['Frontend/Dockerfile',                                    'Multi-stage: Node 20 builds Angular, Nginx Alpine serves dist/'],
        ['Frontend/nginx.conf',                                    'Nginx config — serves SPA, handles HTML5 pushState routing'],
    ],
    col_widths=[3.2, 3.3]
)
doc.add_page_break()

heading('4.3 Microservices Structure', 2)
add_table(
    headers=['Service', 'Port', 'Key Files', 'Database'],
    rows=[
        ['ApiGateway',        '8080', 'Program.cs, appsettings.json (YARP routes), HttpVersionTransform.cs, Dockerfile', 'None'],
        ['AuthService',       '5001', 'AppUser.cs, AuthController.cs, UsersController.cs, JwtService.cs, AuthDtos.cs, Data/AuthDbContext.cs, Data/DbSeeder.cs', 'FreshMart_Auth'],
        ['ProductService',    '5002', 'Product.cs, Category.cs, Review.cs, OrderProjection.cs, ProductsController.cs, CategoriesController.cs, ReviewsController.cs, Data/ProductDbContext.cs, Data/ProductSeeder.cs', 'FreshMart_Product'],
        ['OrderService',      '5003', 'Cart.cs, Order.cs, Coupon.cs, CartController.cs, OrdersController.cs, CouponsController.cs, Services/PaymentServiceClient.cs, Services/NotificationService.cs, Data/OrderDbContext.cs, Data/OrderSeeder.cs', 'FreshMart_Order'],
        ['PaymentService',    '5004', 'Models/Payment.cs, IPaymentService.cs, PaymentController.cs, PaymentDtos.cs, Services/PaymentService.cs, Services/OrderServiceClient.cs, Data/PaymentDbContext.cs', 'FreshMart_Payment'],
        ['NotificationService','5005','Notification.cs, NotificationHub.cs, NotificationsController.cs, InternalNotificationController.cs, NotificationService.cs, EmailService.cs, Data/NotificationDbContext.cs', 'FreshMart_Notification'],
        ['SupportService',    '5006', 'SupportTicket.cs, SupportMessage.cs, SupportHub.cs, SupportController.cs, Data/SupportDbContext.cs', 'FreshMart_Support'],
        ['SharedModels',      'N/A',  'Events.cs (OrderPlacedEvent, OrderStatusChangedEvent, PaymentCompletedEvent), SharedModels.csproj (MassTransit 8.3.6)', 'None'],
    ],
    col_widths=[1.4, 0.5, 3.5, 1.1]
)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. FLOWCHARTS
# ══════════════════════════════════════════════════════════════════════════════
heading('5. Process Flowcharts', 1)

heading('5.1 Checkout & Payment Flow', 2)
body('The checkout flow is the most complex in FreshMart. It spans three microservices: OrderService creates the order, PaymentService creates the Razorpay order and verifies payment, and NotificationService sends the confirmation. The flowchart below shows every decision point and service call.')
insert_image('10_checkout_flow.png', width_inches=4.5,
             caption='Figure 4: Checkout & Payment Flowchart — OrderService + PaymentService + NotificationService')

heading('5.2 Order Status Lifecycle', 2)
body('Every order passes through a defined status lifecycle. Each status change is triggered by a staff member (Admin, StoreManager, or DeliveryDriver) and automatically triggers a SignalR real-time notification and a transactional email to the customer via NotificationService.')
insert_image('11_order_status_flow.png', width_inches=6.8,
             caption='Figure 5: Order Status Lifecycle — Status transitions, who can update, and side effects')
doc.add_page_break()

heading('5.3 Request Flow Summary', 2)
body('All requests from the browser go through the same path: Browser -> Nginx -> Angular -> API Gateway -> Microservice -> SQL Server. The table below summarizes the key flows.')
add_table(
    headers=['Flow', 'Services Involved', 'Key Steps'],
    rows=[
        ['User Login',          'AuthService',
         'POST /auth/login -> BCrypt verify -> JWT issued -> RefreshToken saved to DB'],
        ['Product Search',      'ProductService',
         'GET /products?query= -> EF Core LINQ search -> DiscountedPrice computed -> PaginatedResult returned'],
        ['Add to Cart',         'OrderService',
         'POST /cart/items -> Stock validated from local cache -> CartItem saved -> SubTotal + isOverBudget returned'],
        ['Place Order',         'OrderService -> PaymentService',
         'POST /orders -> Coupon validated -> Order created -> PaymentService called -> Razorpay order created -> razorpayOrderId returned'],
        ['Verify Payment',      'PaymentService',
         'POST /payment/verify -> HMAC SHA256 verified -> Payment.Status = Paid'],
        ['Complete Payment',    'OrderService -> NotificationService (HTTP + RabbitMQ)',
         'POST /orders/{id}/complete-payment -> Cart cleared -> OrderPlacedEvent published to RabbitMQ -> ProductService decrements stock -> NotificationService sends push + email'],
        ['Update Order Status', 'OrderService -> RabbitMQ -> NotificationService',
         'PATCH /orders/{id}/status -> Status updated -> OrderStatusChangedEvent published -> NotificationService sends push + email to customer'],
        ['Support Chat',        'SupportService (SignalR)',
         'POST /support/tickets/{id}/messages -> Message saved -> SignalR broadcast to ticket:{id} group'],
        ['Real-Time Notif',     'NotificationService (SignalR)',
         'POST /internal/user -> Notification saved to DB -> SignalR send to user:{userId} group'],
    ],
    col_widths=[1.4, 1.8, 3.3]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. ER DIAGRAMS
# ══════════════════════════════════════════════════════════════════════════════
heading('6. Entity-Relationship Diagrams', 1)
body('Each microservice has its own isolated SQL Server database. There are no cross-database foreign keys. Relationships between services are maintained via Guid references only. The yellow highlighted rows in each ER diagram indicate Primary Keys (PK). Gold key icon = PK field.')

heading('6.1 AuthService — FreshMart_Auth', 2)
body('Single entity database. All user identity data lives here. AppUser.Id is referenced as a plain Guid by all other services. BCrypt hashes passwords. RefreshToken enables 7-day session persistence. GoogleId enables social login.')
insert_image('04_er_auth.png', width_inches=5.0,
             caption='Figure 6: ER Diagram — AuthService (FreshMart_Auth Database)')

heading('6.2 ProductService — FreshMart_Product', 2)
body('Three main entities: Category (self-referencing hierarchy), Product (belongs to Category), and Review (belongs to Product). OrderProjection and AppUser are local read-only projections from other services — they have no enforced FK to external databases.')
insert_image('05_er_product.png', width_inches=6.8,
             caption='Figure 7: ER Diagram — ProductService (FreshMart_Product Database)')
doc.add_page_break()

heading('6.3 OrderService — FreshMart_Order', 2)
body('Most complex database. Cart has CartItems. Order has OrderItems. Coupon is standalone. A local Product cache table stores price and discount data synced from ProductService to avoid cross-service calls during checkout. ProductName is snapshotted in OrderItem to preserve order history.')
insert_image('06_er_order.png', width_inches=6.8,
             caption='Figure 8: ER Diagram — OrderService (FreshMart_Order Database)')

heading('6.4 PaymentService — FreshMart_Payment', 2)
body('Single entity database. Payment stores all Razorpay transaction data. Status is an enum (0=Pending, 1=Paid, 2=Failed, 3=Refunded, 4=Cancelled). UserId and OrderId are plain Guid references — no enforced FK to other service databases.')
insert_image('07_er_payment.png', width_inches=5.0,
             caption='Figure 9: ER Diagram — PaymentService (FreshMart_Payment Database)')
doc.add_page_break()

heading('6.5 NotificationService — FreshMart_Notification', 2)
body('Notification entity stores all in-app notifications. AppUser is a local projection for user targeting. The NotificationHub uses UserId to route SignalR messages to the correct connected client via group user:{userId}.')
insert_image('08_er_notification.png', width_inches=6.5,
             caption='Figure 10: ER Diagram — NotificationService (FreshMart_Notification Database)')

heading('6.6 SupportService — FreshMart_Support', 2)
body('SupportTicket has many SupportMessages. Status auto-transitions from Open to InProgress on first staff reply (IsStaff=true). ResolvedAt is set when status becomes Resolved or Closed. SupportHub broadcasts messages in real time to all clients in group ticket:{ticketId}.')
insert_image('09_er_support.png', width_inches=6.5,
             caption='Figure 11: ER Diagram — SupportService (FreshMart_Support Database)')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. INTER-SERVICE COMMUNICATION
# ══════════════════════════════════════════════════════════════════════════════
heading('7. Inter-Service Communication', 1)
body('FreshMart uses two communication patterns: synchronous HTTP (typed HttpClient) for critical real-time calls, and asynchronous RabbitMQ events via MassTransit for decoupled background processing. Services call each other directly using Docker internal DNS (container names as hostnames).')

heading('7.1 Synchronous HTTP Calls', 2)
add_table(
    headers=['Caller', 'Called Service', 'Endpoint', 'When', 'Critical?'],
    rows=[
        ['OrderService',  'PaymentService',      'POST /api/v1/payment/create-order',                    'Customer places order',              'YES — order fails if payment service down'],
        ['OrderService',  'NotificationService', 'POST /api/v1/notifications/internal/user',             'Payment completed',                  'NO — wrapped in try/catch'],
        ['OrderService',  'NotificationService', 'POST /api/v1/notifications/internal/role',             'New order placed',                   'NO — wrapped in try/catch'],
        ['OrderService',  'NotificationService', 'POST /api/v1/notifications/internal/email/order-placed','Payment completed',                 'NO — wrapped in try/catch'],
        ['OrderService',  'NotificationService', 'POST /api/v1/notifications/internal/email/order-status','Order status updated',              'NO — wrapped in try/catch'],
        ['OrderService',  'AuthService',         'GET /api/v1/users/{customerId}',                       'Order status updated (for email)',    'NO — wrapped in try/catch'],
        ['PaymentService','OrderService',         'GET /api/v1/orders/{orderId}',                         'Payment creation (order lookup)',     'NO — optional enrichment'],
    ],
    col_widths=[1.2, 1.3, 2.3, 1.6, 1.1]
)

heading('7.2 Async RabbitMQ Events (MassTransit)', 2)
body('RabbitMQ is used for decoupled async communication. Events are defined in SharedModels/Events.cs and consumed by multiple services. MassTransit retry policy: 1s -> 5s -> 15s before dead-lettering.')
add_table(
    headers=['Publisher', 'Event', 'Consumer(s)', 'Effect'],
    rows=[
        ['OrderService',   'OrderPlacedEvent',         'ProductService',      'Decrements stock for each ordered item in a DB transaction'],
        ['OrderService',   'OrderPlacedEvent',         'NotificationService', 'Sends confirmation email + in-app push notification to customer'],
        ['OrderService',   'OrderStatusChangedEvent',  'NotificationService', 'Sends status update email + in-app push notification to customer'],
        ['PaymentService', 'PaymentCompletedEvent',    'OrderService',        'Transitions order status from Pending to Processing automatically'],
    ],
    col_widths=[1.3, 1.8, 1.5, 2.9]
)

heading('7.3 API Gateway Routing Table', 2)
add_table(
    headers=['Route Name', 'URL Pattern', 'Target Service', 'Port'],
    rows=[
        ['auth-route',          '/api/v1/auth/**',          'AuthService',         '5001'],
        ['users-route',         '/api/v1/users/**',         'AuthService',         '5001'],
        ['products-route',      '/api/v1/products/**',      'ProductService',      '5002'],
        ['categories-route',    '/api/v1/categories/**',    'ProductService',      '5002'],
        ['orders-route',        '/api/v1/orders/**',        'OrderService',        '5003'],
        ['cart-route',          '/api/v1/cart/**',          'OrderService',        '5003'],
        ['coupons-route',       '/api/v1/coupons/**',       'OrderService',        '5003'],
        ['payment-route',       '/api/v1/payment/**',       'PaymentService',      '5004'],
        ['notifications-route', '/api/v1/notifications/**', 'NotificationService', '5005'],
        ['notif-hub-route',     '/hubs/notifications/**',   'NotificationService', '5005'],
        ['support-route',       '/api/v1/support/**',       'SupportService',      '5006'],
        ['support-hub-route',   '/hubs/support/**',         'SupportService',      '5006'],
    ],
    col_widths=[1.5, 1.8, 1.7, 0.7]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. DESIGN PATTERNS
# ══════════════════════════════════════════════════════════════════════════════
heading('8. Design Patterns Used', 1)
add_table(
    headers=['Pattern', 'Where Used', 'File(s)', 'Description'],
    rows=[
        ['API Gateway',          'Architecture level',    'ApiGateway/Program.cs, appsettings.json',
         'Single entry point. YARP routes to correct service. JWT validated once centrally.'],
        ['Dependency Injection', 'All services',          'Every Program.cs',
         'ASP.NET Core DI. Services registered as Scoped. Injected via constructor parameters.'],
        ['Interface Abstraction','PaymentService',        'IPaymentService.cs + Services/PaymentService.cs',
         'IPaymentService defines contract; concrete class implements it. Enables mocking in tests.'],
        ['Repository Pattern',   'All services',          'Data/*DbContext.cs files',
         'EF Core DbContext acts as repository. Business logic queries through DbContext, not raw SQL.'],
        ['Service Client',       'OrderService',          'Services/PaymentServiceClient.cs, Services/NotificationService.cs',
         'Typed HttpClient wrappers for inter-service HTTP calls with configured base addresses.'],
        ['Seeder Pattern',       'Auth, Product, Order',  'Data/DbSeeder.cs, Data/ProductSeeder.cs, Data/OrderSeeder.cs',
         'Static SeedAsync() called on startup to populate initial data (users, products, coupons).'],
        ['Middleware Pipeline',  'All services',          'Every Program.cs',
         'UseCors -> UseAuthentication -> UseAuthorization -> MapControllers. Order is critical.'],
        ['Snapshot Pattern',     'OrderService',          'OrderItem.cs (ProductName, UnitPrice fields)',
         'ProductName and UnitPrice snapshotted at order time. Preserves history if product changes.'],
        ['Projection Pattern',   'Product, Notif, Support','OrderProjection.cs, AppUser.cs (per service)',
         'Local read-only copies of data from other services. Avoids cross-service DB calls.'],
        ['Multi-Stage Docker',   'All services + Frontend','Every Dockerfile',
         'Stage 1: SDK/Node builds. Stage 2: Runtime/Nginx serves. Final image has no build tools.'],
        ['JWT Stateless Auth',   'All services',          'JwtService.cs + every Program.cs',
         'JWT signed with shared secret. Any service verifies without calling AuthService.'],
        ['Event-Driven Messaging','OrderService, PaymentService, ProductService, NotificationService',
         'SharedModels/Events.cs, Messaging/MassTransitConfig.cs, Consumers/',
         'MassTransit + RabbitMQ for async events. Publisher fires and forgets. Consumers retry on failure (1s/5s/15s).'],
    ],
    col_widths=[1.4, 1.4, 2.0, 2.7]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. DOCKER COMPOSE & QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
heading('9. Docker Compose & Quick Reference', 1)

heading('9.1 Container Startup Order', 2)
add_table(
    headers=['Order', 'Container', 'Port', 'Depends On', 'Database'],
    rows=[
        ['1st', 'freshmart-sqlserver',     '1433', 'None (healthcheck: sqlcmd SELECT 1)',                  'N/A — hosts all 6 DBs'],
        ['2nd', 'freshmart-auth',          '5001', 'sqlserver (healthy)',                                  'FreshMart_Auth'],
        ['2nd', 'freshmart-products',      '5002', 'sqlserver (healthy)',                                  'FreshMart_Product'],
        ['2nd', 'freshmart-orders',        '5003', 'sqlserver (healthy)',                                  'FreshMart_Order'],
        ['2nd', 'freshmart-payment',       '5004', 'sqlserver (healthy)',                                  'FreshMart_Payment'],
        ['2nd', 'freshmart-notifications', '5005', 'sqlserver (healthy)',                                  'FreshMart_Notification'],
        ['2nd', 'freshmart-support',       '5006', 'sqlserver (healthy)',                                  'FreshMart_Support'],
        ['3rd', 'freshmart-gateway',       '8080', 'auth, products, orders, payment, notifications, support', 'None'],
        ['4th', 'freshmart-frontend',      '80',   'api-gateway',                                          'None'],
    ],
    col_widths=[0.5, 1.7, 0.5, 2.5, 1.3]
)

heading('9.2 Run Commands', 2)
body('From the workspace root (where docker-compose.yml is):')
code_block('docker compose up --build          # Build and start all containers')
code_block('docker compose up --build -d       # Run in background (detached)')
code_block('docker compose down               # Stop all containers')
code_block('docker compose down -v            # Stop and wipe all database data')
code_block('docker compose logs -f            # Stream all logs')
code_block('docker compose logs -f auth-service  # Stream logs for one service')

heading('9.3 Access URLs', 2)
add_table(
    headers=['Service', 'URL', 'Notes'],
    rows=[
        ['Frontend (Angular)',    'http://localhost',              'Main application UI'],
        ['API Gateway',          'http://localhost:8080',         'All API calls go through here'],
        ['AuthService Swagger',  'http://localhost:5001/swagger', 'Auth API documentation'],
        ['ProductService Swagger','http://localhost:5002/swagger', 'Product API documentation'],
        ['OrderService Swagger',  'http://localhost:5003/swagger', 'Order/Cart API documentation'],
        ['PaymentService Swagger','http://localhost:5004/swagger', 'Payment API documentation'],
        ['NotifService Swagger',  'http://localhost:5005/swagger', 'Notification API documentation'],
        ['SupportService Swagger','http://localhost:5006/swagger', 'Support API documentation'],
    ],
    col_widths=[1.8, 2.2, 2.5]
)

heading('9.4 Test Accounts (Seeded on Startup)', 2)
add_table(
    headers=['Role', 'Email', 'Password', 'Access'],
    rows=[
        ['Admin',          'admin@grocery.com',    'Admin@123',    'Full platform access — all routes'],
        ['Store Manager',  'manager@grocery.com',  'Manager@123',  'Products, orders, support management'],
        ['Delivery Driver','driver@grocery.com',   'Driver@123',   'Delivery orders only (/delivery)'],
        ['Customer',       'customer@grocery.com', 'Customer@123', 'Shopping, orders, support, reviews'],
    ],
    col_widths=[1.3, 2.0, 1.3, 2.0]
)

heading('9.5 Key Business Rules', 2)
add_table(
    headers=['Rule', 'Value', 'Enforced In'],
    rows=[
        ['Delivery Fee',         'Rs.0 if SubTotal >= Rs.500, else Rs.49',    'OrderService — CreateOrder'],
        ['Tax Rate',             '5% of SubTotal',                            'OrderService — CreateOrder'],
        ['JWT Expiry',           '1 hour (access token)',                     'AuthService — JwtService'],
        ['Refresh Token Expiry', '7 days',                                   'AuthService — AuthController'],
        ['Password Hashing',     'BCrypt cost factor 12',                    'AuthService — AuthController'],
        ['Low Stock Threshold',  'StockQuantity < 10',                       'ProductService — ProductsController'],
        ['Review Limit',         'One review per customer per product',      'ProductService — ReviewsController'],
        ['Notification Limit',   'Last 50 returned per user',               'NotificationService — NotificationsController'],
        ['Coupon Validation',    'IsActive + not expired + UsedCount < UsageLimit + SubTotal >= MinOrderAmount', 'OrderService — OrdersController'],
        ['Order Status Flow',    'Pending -> Processing -> Shipped -> OutForDelivery -> Delivered (or Cancelled)', 'OrderService — OrdersController'],
        ['Ticket Status Flow',   'Open -> InProgress (first staff reply) -> Resolved -> Closed', 'SupportService — SupportController'],
        ['DiscountedPrice',      'Price x (1 - DiscountPercent / 100)',      'ProductService — ProductsController.ToDto()'],
    ],
    col_widths=[1.8, 2.2, 2.5]
)

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run('End of Document — FreshMart Project Structure & Flow Document')
r.bold = True
r.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

doc.save('FreshMart_ProjectStructure.docx')
print('SUCCESS: FreshMart_ProjectStructure.docx has been generated!')
